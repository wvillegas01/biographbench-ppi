"""Surgically revise the working manuscript to match Phase 1 PPI results."""

from __future__ import annotations

import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent
INPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript.docx"
OUTPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised.docx"
BACKUP_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript.before_phase1_ppi_revision.docx"

CLASSICAL = ROOT / "results" / "phase1" / "ppi_link_prediction_baselines.csv"
NODE2VEC = ROOT / "results" / "phase1" / "ppi_node2vec_link_prediction.csv"
FIG_DIR = ROOT / "Articulo" / "figures"

DATASET_LABELS = {
    "string_human_physical_v12": "STRING",
    "biogrid_human_physical": "BioGRID",
    "string_human_physical_no_biogrid_overlap": "STRING without BioGRID",
    "biogrid_human_physical_no_string_overlap": "BioGRID without STRING",
}

DATASET_ORDER = [
    "string_human_physical_v12",
    "biogrid_human_physical",
    "string_human_physical_no_biogrid_overlap",
    "biogrid_human_physical_no_string_overlap",
]

NEGATIVE_ORDER = ["random", "degree_matched", "two_hop"]
NEGATIVE_LABELS = {
    "random": "random",
    "degree_matched": "degree-matched",
    "two_hop": "two-hop",
}

MODEL_ORDER = ["logistic_regression", "random_forest", "hist_gradient_boosting", "node2vec_walk_logreg"]
MODEL_LABELS = {
    "logistic_regression": "LogReg",
    "random_forest": "RF",
    "hist_gradient_boosting": "HGB",
    "node2vec_walk_logreg": "node2vec",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_table(table) -> None:
    element = table._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_paragraph_text(paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def delete_between_headings(doc: Document, start_heading: str, end_heading: str) -> None:
    start = find_paragraph(doc, start_heading)
    end = find_paragraph(doc, end_heading)
    deleting = False
    for paragraph in list(doc.paragraphs):
        if paragraph._p is start._p:
            deleting = True
            continue
        if paragraph._p is end._p:
            break
        if deleting:
            delete_paragraph(paragraph)


def insert_picture_after(paragraph, image_path: Path, width_inches: float = 6.5):
    pic = insert_paragraph_after(paragraph)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return pic


def read_results() -> pd.DataFrame:
    classical = pd.read_csv(CLASSICAL)
    node = pd.read_csv(NODE2VEC)
    node["model_family"] = "embedding"
    node["train_seconds"] = node["walk_seconds"] + node["embedding_seconds"] + node["decoder_seconds"]
    shared = [
        "dataset",
        "negative_strategy",
        "seed",
        "split",
        "model",
        "model_family",
        "auroc",
        "auprc",
        "brier",
        "nll",
        "ece_10",
        "train_seconds",
        "inference_seconds",
        "num_train",
        "num_eval",
    ]
    return pd.concat([classical[shared], node[shared]], ignore_index=True)


def mean_ci(values: pd.Series) -> tuple[float, float, float, float]:
    values = values.astype(float)
    mean = float(values.mean())
    sd = float(values.std(ddof=1))
    ci = 1.96 * sd / math.sqrt(values.count())
    return mean, sd, mean - ci, mean + ci


def table_set_size(table, rows: int, cols: int) -> None:
    while len(table.columns) < cols:
        table.add_column(Inches(1.0))
    while len(table.rows) < rows:
        table.add_row()
    while len(table.rows) > rows:
        table._tbl.remove(table.rows[-1]._tr)


def fill_cell(cell, text: str) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = None


def update_front_matter(doc: Document) -> None:
    set_paragraph_text(
        doc.paragraphs[2],
        "BioGraphBench-PPI: audit-first evaluation of negative-sampling contracts and structural baselines for protein interaction link prediction",
        "Title",
    )
    set_paragraph_text(doc.paragraphs[3], "Running title: Audit-first PPI link prediction")
    set_paragraph_text(doc.paragraphs[7], "Manuscript counts: [update before submission]; 4 figures; 4 tables; 25 references.")
    set_paragraph_text(
        doc.paragraphs[8],
        "WORKING-MANUSCRIPT NOTE — remove before submission: this revision aligns the manuscript with the completed Phase 1 PPI analysis: 10 seeds, three negative-sampling regimes, classical structural baselines, a node2vec-compatible random-walk baseline, calibration, runtime diagnostics, and paired statistical summaries.",
    )


def update_abstract_intro(doc: Document) -> None:
    abstract = (
        "Introduction: Biomedical graph-learning results are difficult to compare when source versions, identifier rules, data splits, negative sampling, and structural leakage controls are incompletely documented. "
        "Methods: We revised BioGraphBench as an audit-first PPI link-prediction benchmark using STRING v12.0, BioGRID 5.0.260, and two cross-resource no-overlap ablations. Each task was evaluated over 10 predefined seeds under three negative-sampling contracts: random non-edges, degree-matched non-edges, and two-hop hard non-edges sampled from the training graph. We compared local structural heuristics, logistic regression, random forest, histogram gradient boosting, and an auditable node2vec-compatible random-walk baseline, reporting AUROC, AUPRC, calibration, runtime, confidence intervals, paired Wilcoxon tests, and effect sizes. "
        "Results: Random negatives produced high HGB AUPRC values (0.943-0.964), but degree-matched negatives reduced AUPRC by 0.049-0.259 depending on dataset. Two-hop negatives further changed the difficulty profile and model ranking, with HGB/RF AUPRC values of 0.774-0.857 and node2vec AUPRC values of 0.605-0.736. "
        "Discussion: These results show that PPI benchmark conclusions depend strongly on the negative-sampling contract and graph source. BioGraphBench-PPI therefore contributes an auditable evaluation protocol rather than a claim that PPI prediction is solved."
    )
    set_paragraph_text(doc.paragraphs[10], abstract)
    set_paragraph_text(
        doc.paragraphs[11],
        "Keywords: protein-protein interactions, link prediction, negative sampling, node2vec, biomedical networks, benchmarking, reproducibility, calibration",
    )

    replacements = {
        "We present BioGraphBench, an audit-first benchmark": (
            "We present BioGraphBench-PPI, a focused audit-first benchmark for reproducible protein-interaction link prediction. "
            "The central research question is whether open PPI resources can be transformed into link-prediction tasks whose evidence chain is reconstructible and whose conclusions remain stable under explicit negative-sampling contracts."
        ),
        "The contribution is both methodological and empirical.": (
            "The contribution is both methodological and empirical. Methodologically, BioGraphBench-PPI makes source provenance, split contracts, negative sampling, leakage control, and feature construction first-class experimental objects. "
            "Empirically, it establishes a multi-seed structural baseline for four PPI tasks and shows that random, degree-matched, and two-hop negatives lead to materially different performance estimates and model rankings."
        ),
    }
    for paragraph in doc.paragraphs:
        for prefix, text in replacements.items():
            if paragraph.text.startswith(prefix):
                set_paragraph_text(paragraph, text)


def update_methods(doc: Document) -> None:
    replacements = {
        "where Gₖ is the canonical graph": (
            "where Gk is the canonical graph, Tk is the task definition, Sk contains the predefined multi-seed data partitions, Phik specifies feature construction, and Mk defines the evaluation metrics. "
            "An experiment is accepted only when its source, transformation rules, split strategy, negative-sampling contract, features, model configuration, seed, and outputs have machine-readable records."
        ),
        "Five tasks were accepted": (
            "Four PPI link-prediction tasks were retained for the focused Phase 1 manuscript (Table 1). For each graph, positive edges were shuffled with 10 predefined seeds (42-51) and partitioned into 80% training, 10% validation, and 10% testing subsets while preserving a spanning forest in the training graph."
        ),
        "All structural features were computed": (
            "All structural features were computed from Gtrain = (V, E+train) only. For each seed, three balanced negative-sampling contracts were generated. The random contract samples uniformly from unordered node pairs absent from the complete observed positive-edge set. The degree-matched contract approximates the endpoint log-degree-bin distribution of the corresponding positive partition. The two-hop contract samples non-edges whose endpoints share at least one neighbor in the training-positive graph, making the negative examples structurally closer without using validation or test positives to define hardness."
        ),
        "Each individual heuristic served": (
            "Each individual heuristic served as an unsupervised score. The complete feature vector was supplied to logistic regression, random forest, and histogram gradient boosting (HGB) as supervised structural baselines. We also included an auditable node2vec-compatible random-walk baseline: unbiased walks (p = q = 1) were used to learn embeddings, and a logistic edge decoder was trained on dot product, cosine similarity, L1 distance, and L2 distance features."
        ),
        "The software environment was pinned": (
            "The software environment was pinned to NumPy 1.24.3, pandas 2.0.3, SciPy 1.10.1, scikit-learn 1.3.2, and PyTorch 2.4.1+CPU. Data checks, split validation, feature-policy tests, artifact-presence tests, raw per-seed outputs, statistical summaries, and regenerated figure/table scripts were retained with the benchmark. All PPI results in the present manuscript come from the 10-seed Phase 1 protocol."
        ),
    }
    for paragraph in list(doc.paragraphs):
        for prefix, text in replacements.items():
            if paragraph.text.startswith(prefix):
                set_paragraph_text(paragraph, text)

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == "Multilabel node-classification baselines":
            delete_paragraph(paragraph)
        elif paragraph.text.startswith("The node-classification input was"):
            delete_paragraph(paragraph)
        elif paragraph.text.startswith("The secondary task used OBNB"):
            delete_paragraph(paragraph)

    for paragraph in list(doc.paragraphs):
        if "w:drawing" in paragraph._p.xml and paragraph._p.getnext() is not None:
            next_text = ""
            for candidate in doc.paragraphs:
                if candidate._p is paragraph._p.getnext():
                    next_text = candidate.text
                    break
            if next_text.startswith("Figure 1 | Audit-first"):
                delete_paragraph(paragraph)
                break
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith("Figure 1 | Audit-first"):
            delete_paragraph(paragraph)


def rebuild_results(doc: Document) -> None:
    delete_between_headings(doc, "Results", "Discussion")
    anchor = find_paragraph(doc, "Results")

    blocks = [
        ("Four auditable PPI tasks and 120 split contracts", "Heading 2"),
        (
            "BioGraphBench-PPI retained four link-prediction tasks whose source files, canonicalization rules, split manifests, and output records could be reconstructed (Tables 1 and 2). The two complete PPI graphs contained 18,767-20,376 nodes and 738,805-961,531 canonical undirected edges. Removing cross-resource overlap reduced BioGRID to 538,003 edges and STRING to 314,539 edges, providing ablation tasks that test whether performance is driven by shared database content. Across the four datasets, three negative-sampling contracts, and 10 seeds, the benchmark produced 120 split manifests, 1,680 classical-baseline result rows, and 240 node2vec-compatible result rows.",
            None,
        ),
        ("Random negatives inflate apparent link-prediction performance", "Heading 2"),
        (
            "Under uniformly sampled random negatives, HGB achieved high mean test AUPRC across all PPI tasks: 0.950 for STRING, 0.943 for BioGRID, 0.964 for STRING without BioGRID, and 0.942 for BioGRID without STRING. The seed-to-seed dispersion was small, with SD values below 0.001 in the HGB summaries, showing that the random-negative protocol is reproducible but structurally easy. Figure 1 shows these tight distributions and makes clear that high random-negative performance should be interpreted as a controlled structural baseline rather than evidence of prospective biological completion.",
            None,
        ),
    ]
    for text, style in blocks:
        anchor = insert_paragraph_after(anchor, text, style)
    anchor = insert_picture_after(anchor, FIG_DIR / "fig1_hgb_auprc_seed_distributions.png")
    anchor = insert_paragraph_after(
        anchor,
        "Figure 1 | HGB performance distributions across 10 seeds. Violin densities and seed-level points summarize test AUPRC for complete PPI graphs and cross-database no-overlap ablations under random, degree-matched, and two-hop negative-sampling contracts.",
        "Caption",
    )

    for text, style in [
        (
            "Degree-matched and two-hop negatives materially change the task",
            "Heading 2",
        ),
        (
            "Replacing random negatives with degree-matched negatives reduced HGB AUPRC in every dataset. The absolute reductions were 0.069 for STRING, 0.172 for BioGRID, 0.049 for STRING without BioGRID, and 0.259 for BioGRID without STRING. The two-hop contract produced a different hard-negative profile: HGB AUPRC was 0.856 for STRING, 0.774 for BioGRID, 0.847 for STRING without BioGRID, and 0.849 for BioGRID without STRING. Thus, negative difficulty is not a single monotonic property; it depends on whether the control preserves degree distribution, local closure, or both.",
            None,
        ),
    ]:
        anchor = insert_paragraph_after(anchor, text, style)
    anchor = insert_picture_after(anchor, FIG_DIR / "fig2_hgb_negative_regime_drops.png")
    anchor = insert_paragraph_after(
        anchor,
        "Figure 2 | Seed-paired performance change induced by harder negative sampling. Points show per-seed AUPRC differences between random negatives and the harder contracts; whiskers show 95% confidence intervals across 10 seeds.",
        "Caption",
    )

    for text, style in [
        ("Model rankings depend on the negative-sampling contract", "Heading 2"),
        (
            "The expanded baseline set changes the interpretation of model strength (Table 3). HGB and random forest were generally strongest in complete STRING and BioGRID settings, but the model-family profiles were not invariant. In BioGRID without STRING, node2vec-compatible embeddings exceeded HGB/RF under degree-matched negatives (AUPRC 0.725 versus 0.684-0.687), whereas HGB recovered the highest value under two-hop negatives (0.849). In STRING without BioGRID, random forest slightly exceeded HGB under two-hop negatives (0.851 versus 0.847). These shifts support the benchmark thesis: model claims must be attached to a precise negative-sampling contract.",
            None,
        ),
    ]:
        anchor = insert_paragraph_after(anchor, text, style)
    anchor = insert_picture_after(anchor, FIG_DIR / "fig3_model_family_regime_profiles.png")
    anchor = insert_paragraph_after(
        anchor,
        "Figure 3 | Model-family profiles across negative-sampling regimes. Mean test AUPRC across 10 seeds is shown for logistic regression, random forest, HGB, and the node2vec-compatible random-walk baseline in each PPI task.",
        "Caption",
    )

    for text, style in [
        ("Calibration and cost expose practical differences among baselines", "Heading 2"),
        (
            "Ranking metrics alone did not fully characterize the baselines. HGB was consistently well calibrated relative to the other model families, with ECE-10 values generally below the broader logistic-regression and node2vec ranges (Table 4; Figure 4A). Runtime diagnostics also separated practical baselines from purely predictive ones: logistic regression was fastest, random forest was most expensive, and HGB/node2vec occupied an intermediate range. These diagnostics are important because a benchmark baseline should be not only accurate, but also reproducible, calibratable, and computationally interpretable.",
            None,
        ),
    ]:
        anchor = insert_paragraph_after(anchor, text, style)
    anchor = insert_picture_after(anchor, FIG_DIR / "fig4_calibration_and_runtime.png")
    anchor = insert_paragraph_after(
        anchor,
        "Figure 4 | Calibration and computational-cost diagnostics. (A) ECE-10 distributions by model family and negative-sampling contract. (B) Training-time distributions by model family, with points showing dataset-level runs.",
        "Caption",
    )

    for text, style in [
        ("Empirical requirements for future graph models", "Heading 2"),
        (
            "The Phase 1 results establish a more demanding entry point for future GNN submissions. A new PPI link-prediction model should be compared against HGB, random forest, and node2vec-compatible embeddings; should report mean, SD, confidence intervals, paired tests, and calibration; and should demonstrate that any gain persists under random, degree-matched, and two-hop negative contracts. Without these controls, apparent improvements may reflect an easier negative set rather than better biological graph learning.",
            None,
        ),
    ]:
        anchor = insert_paragraph_after(anchor, text, style)


def rebuild_discussion(doc: Document) -> None:
    delete_between_headings(doc, "Discussion", "Conclusion")
    anchor = find_paragraph(doc, "Discussion")
    paragraphs = [
        "BioGraphBench-PPI shows that the scientific value of a PPI link-prediction benchmark depends less on reporting a high score than on making the experimental contract auditable. The same biological graph can support very different conclusions depending on whether negatives are sampled uniformly, matched by endpoint degree, or constrained to lie at distance two in the training graph. This result directly addresses a recurring weakness in graph-learning evaluations: performance is often attributed to model architecture even when the negative set encodes much of the task difficulty.",
        "The multi-seed results also change the status of the benchmark. The earlier pilot analysis was reproducible but did not estimate sampling variance. The revised protocol shows that the main findings are stable across 10 seeds while preserving enough resolution to support confidence intervals, Wilcoxon tests, and paired effect sizes. This strengthens the article because reproducibility is no longer only a software property; it is part of the statistical evidence.",
        "The model-family comparison provides a useful caution. HGB and random forest are strong structural baselines, but node2vec-compatible embeddings are not uniformly weak, and in one degree-matched no-overlap setting they outperform the supervised pair-feature models. Conversely, two-hop negatives penalize node2vec more strongly in several datasets and can restore HGB/RF dominance. These crossovers indicate that model rankings are conditional statements, not general truths about PPI prediction.",
        "The findings are consistent with prior work showing that graph structure can be highly predictive in molecular-interaction tasks, but they sharpen the methodological implication: stronger architectures should be evaluated against negative controls that remove simple degree and closure shortcuts. The benchmark therefore complements GNN-oriented PPI studies by providing a transparent non-GNN floor that future graph neural models must exceed under identical data contracts.",
        "Several limitations remain. First, the current manuscript is deliberately focused on PPI link prediction and should not be presented as a general network-bioinformatics benchmark. Second, the two-hop strategy is a hard-negative proxy based on local closure; it does not exhaust community-aware, biology-aware, or temporally held-out alternatives. Third, the node2vec-compatible baseline uses unbiased walks and a logistic decoder rather than an exhaustive hyperparameter search. Finally, GCN, GraphSAGE, GAT/GIN, and SEAL-style link-prediction models remain necessary for the next phase.",
        "Even with these limitations, the revised study is substantially more defensible than the initial MVP. It now contains multi-seed uncertainty, three negative-sampling regimes, competitive non-GNN and embedding baselines, calibration, runtime diagnostics, and audit records that connect every result back to source data and split manifests.",
    ]
    for text in paragraphs:
        anchor = insert_paragraph_after(anchor, text)


def update_conclusion_and_availability(doc: Document) -> None:
    conclusion = find_paragraph(doc, "BioGraphBench converts open biomedical networks")
    set_paragraph_text(
        conclusion,
        "BioGraphBench-PPI converts open protein-interaction resources into reconstructible link-prediction tasks and shows why negative-sampling contracts must be treated as first-class benchmark objects. Across four PPI datasets, 10 seeds, and three negative regimes, random negatives produced high and stable performance, whereas degree-matched and two-hop negatives changed both absolute AUPRC and model ranking. The strongest practical baselines were HGB and random forest in most settings, but node2vec-compatible embeddings were competitive under selected harder contracts. These findings support a focused article: BioGraphBench-PPI is not a claim that PPI prediction is solved, but a reproducible evidence framework for testing whether future graph models improve beyond structural shortcuts.",
    )
    data_avail = find_paragraph(doc, "BioGraphBench was developed from openly accessible")
    set_paragraph_text(
        data_avail,
        "BioGraphBench-PPI was developed from openly accessible STRING and BioGRID resources. The working project snapshot contains source manifests, canonicalized PPI task definitions, 120 split manifests, raw multi-seed outputs, statistical summaries, regenerated figures, tables, and validation records supporting this manuscript.",
    )


def update_tables(doc: Document, data: pd.DataFrame) -> None:
    # Table 1: accepted PPI tasks only.
    t1 = doc.tables[0]
    table_set_size(t1, 5, 5)
    rows = [
        ["Task", "Source", "Prediction unit", "Primary metrics", "Role"],
        ["STRING physical", "STRING v12.0", "Undirected protein pair", "AUROC, AUPRC", "Main PPI"],
        ["BioGRID physical", "BioGRID 5.0.260", "Undirected protein pair", "AUROC, AUPRC", "Main PPI"],
        ["STRING without BioGRID", "STRING after pair removal", "Undirected protein pair", "AUROC, AUPRC", "Overlap ablation"],
        ["BioGRID without STRING", "BioGRID after pair removal", "Undirected protein pair", "AUROC, AUPRC", "Overlap ablation"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t1.cell(r, c), value)

    # Table 2: PPI scale and per-seed positive partitions.
    t2 = doc.tables[1]
    table_set_size(t2, 5, 6)
    rows = [
        ["Task", "Nodes", "Graph edges", "Train +", "Validation +", "Test +"],
        ["STRING physical", "18,767", "738,805", "591,045", "73,880", "73,880"],
        ["BioGRID physical", "20,376", "961,531", "769,225", "96,153", "96,153"],
        ["STRING without BioGRID", "16,781", "314,539", "251,633", "31,453", "31,453"],
        ["BioGRID without STRING", "19,591", "538,003", "430,403", "53,800", "53,800"],
    ]
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t2.cell(r, c), value)

    test = data[data["split"] == "test"].copy()
    models = test[test["model"].isin(MODEL_ORDER)]
    summary = models.groupby(["dataset", "negative_strategy", "model"])["auprc"].mean().to_dict()

    # Table 3: mean AUPRC by model/regime.
    t3 = doc.tables[2]
    table_set_size(t3, 13, 6)
    rows = [["Task", "Negatives", "LogReg", "RF", "HGB", "node2vec"]]
    for dataset in DATASET_ORDER:
        for neg in NEGATIVE_ORDER:
            rows.append(
                [
                    DATASET_LABELS[dataset],
                    NEGATIVE_LABELS[neg],
                    f"{summary[(dataset, neg, 'logistic_regression')]:.3f}",
                    f"{summary[(dataset, neg, 'random_forest')]:.3f}",
                    f"{summary[(dataset, neg, 'hist_gradient_boosting')]:.3f}",
                    f"{summary[(dataset, neg, 'node2vec_walk_logreg')]:.3f}",
                ]
            )
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t3.cell(r, c), value)

    # Table 4: sensitivity of HGB to negative regimes.
    t4 = doc.tables[3]
    table_set_size(t4, 5, 6)
    rows = [["Task", "Random", "Degree-matched", "Two-hop", "Drop: random-degree", "Drop: random-two-hop"]]
    hgb = test[test["model"] == "hist_gradient_boosting"]
    for dataset in DATASET_ORDER:
        vals = {
            neg: float(hgb[(hgb["dataset"] == dataset) & (hgb["negative_strategy"] == neg)]["auprc"].mean())
            for neg in NEGATIVE_ORDER
        }
        rows.append(
            [
                DATASET_LABELS[dataset],
                f"{vals['random']:.3f}",
                f"{vals['degree_matched']:.3f}",
                f"{vals['two_hop']:.3f}",
                f"{vals['random'] - vals['degree_matched']:.3f}",
                f"{vals['random'] - vals['two_hop']:.3f}",
            ]
        )
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t4.cell(r, c), value)

    # Remove old OBNB table from the PPI-focused manuscript.
    if len(doc.tables) > 4:
        delete_table(doc.tables[4])


def update_table_captions(doc: Document) -> None:
    captions = {
        "Table 1 | Accepted tasks": "Table 1 | Accepted PPI link-prediction tasks in BioGraphBench-PPI.",
        "Table 2 | Scale and fixed partitions": "Table 2 | PPI graph scale and positive partition sizes per seed. Each positive partition is paired with an equal number of negatives under each negative-sampling contract.",
        "Table 3 | Best classical heuristic": "Table 3 | Mean test AUPRC across 10 seeds by dataset, negative-sampling contract, and model family.",
        "Table 4 | Calibration and recorded": "Table 4 | HGB sensitivity to negative-sampling contracts. Drops are seed-averaged AUPRC differences relative to random negatives.",
        "Table 5 | Test performance": "",
    }
    for paragraph in list(doc.paragraphs):
        for prefix, text in captions.items():
            if paragraph.text.startswith(prefix):
                if text:
                    set_paragraph_text(paragraph, text, "Caption")
                else:
                    delete_paragraph(paragraph)


def main() -> int:
    if not BACKUP_DOCX.exists():
        shutil.copy2(INPUT_DOCX, BACKUP_DOCX)

    data = read_results()
    doc = Document(INPUT_DOCX)
    update_front_matter(doc)
    update_abstract_intro(doc)
    update_methods(doc)
    rebuild_results(doc)
    rebuild_discussion(doc)
    update_conclusion_and_availability(doc)
    update_tables(doc, data)
    update_table_captions(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

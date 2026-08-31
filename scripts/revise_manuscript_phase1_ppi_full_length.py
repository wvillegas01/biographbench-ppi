"""Conservative full-length manuscript revision for Phase 1 PPI.

This script starts from the original working manuscript and preserves the long
article architecture: methodological workflow, OBNB secondary section, references
and five original table slots. It only updates claims that became inconsistent
after the completed Phase 1 PPI analysis.
"""

from __future__ import annotations

import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
PROJECT_ROOT = ROOT.parent
INPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript.before_phase1_ppi_revision.docx"
FALLBACK_INPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript.docx"
OUTPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised_full_length.docx"

CLASSICAL = ROOT / "results" / "phase1" / "ppi_link_prediction_baselines.csv"
NODE2VEC = ROOT / "results" / "phase1" / "ppi_node2vec_link_prediction.csv"
FIG_DIR = ROOT / "Articulo" / "figures"

DATASET_LABELS = {
    "string_human_physical_v12": "STRING",
    "biogrid_human_physical": "BioGRID",
    "string_human_physical_no_biogrid_overlap": "STRING without BioGRID",
    "biogrid_human_physical_no_string_overlap": "BioGRID without STRING",
}

DATASET_ORDER = [
    "string_human_physical_v12",
    "biogrid_human_physical",
    "string_human_physical_no_biogrid_overlap",
    "biogrid_human_physical_no_string_overlap",
]

NEGATIVE_ORDER = ["random", "degree_matched", "two_hop"]
NEGATIVE_LABELS = {
    "random": "random",
    "degree_matched": "degree-matched",
    "two_hop": "two-hop",
}

MODEL_ORDER = ["logistic_regression", "random_forest", "hist_gradient_boosting", "node2vec_walk_logreg"]
MODEL_LABELS = {
    "logistic_regression": "LogReg",
    "random_forest": "RF",
    "hist_gradient_boosting": "HGB",
    "node2vec_walk_logreg": "node2vec",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_before(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_text(paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def previous_paragraph(doc: Document, paragraph):
    paragraphs = doc.paragraphs
    for idx, candidate in enumerate(paragraphs):
        if candidate._p is paragraph._p:
            return paragraphs[idx - 1] if idx > 0 else None
    return None


def replace_figure_before_caption(doc: Document, caption_prefix: str, image_path: Path, caption_text: str, width: float = 6.5):
    caption = find_paragraph(doc, caption_prefix)
    old_image = previous_paragraph(doc, caption)
    if old_image is not None and "w:drawing" in old_image._p.xml:
        delete_paragraph(old_image)
    image_para = insert_paragraph_before(caption)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    set_text(caption, caption_text, "Caption")


def add_figure_before_caption(doc: Document, caption_prefix: str, image_path: Path, caption_text: str, width: float = 6.5):
    caption = find_paragraph(doc, caption_prefix)
    image_para = insert_paragraph_before(caption)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    new_caption = insert_paragraph_after(image_para, caption_text, "Caption")
    return new_caption


def add_figure_after_caption(doc: Document, caption_prefix: str, image_path: Path, caption_text: str, width: float = 6.5):
    caption = find_paragraph(doc, caption_prefix)
    image_para = insert_paragraph_after(caption)
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    new_caption = insert_paragraph_after(image_para, caption_text, "Caption")
    return new_caption


def read_results() -> pd.DataFrame:
    classical = pd.read_csv(CLASSICAL)
    node = pd.read_csv(NODE2VEC)
    node["model_family"] = "embedding"
    node["train_seconds"] = node["walk_seconds"] + node["embedding_seconds"] + node["decoder_seconds"]
    shared = [
        "dataset",
        "negative_strategy",
        "seed",
        "split",
        "model",
        "model_family",
        "auroc",
        "auprc",
        "brier",
        "nll",
        "ece_10",
        "train_seconds",
        "inference_seconds",
        "num_train",
        "num_eval",
    ]
    return pd.concat([classical[shared], node[shared]], ignore_index=True)


def table_set_size(table, rows: int, cols: int) -> None:
    while len(table.columns) < cols:
        table.add_column(Inches(1.0))
    while len(table.rows) < rows:
        table.add_row()
    while len(table.rows) > rows:
        table._tbl.remove(table.rows[-1]._tr)


def fill_cell(cell, value: str) -> None:
    cell.text = str(value)


def update_front_and_abstract(doc: Document) -> None:
    set_text(
        doc.paragraphs[2],
        "BioGraphBench-PPI enables audit-first evaluation of negative-sampling contracts and structural baselines in protein interaction networks",
        "Title",
    )
    set_text(doc.paragraphs[3], "Running title: Audit-first PPI link prediction")
    set_text(doc.paragraphs[7], "Manuscript counts: [update before submission]; 6 figures; 5 tables; 25 references.")
    set_text(
        doc.paragraphs[8],
        "WORKING-MANUSCRIPT NOTE — remove before submission: this full-length revision preserves the original article structure while aligning the PPI claims with the completed Phase 1 protocol: 10 seeds, three negative-sampling regimes, classical structural baselines, a node2vec-compatible baseline, statistical summaries, calibration, and runtime diagnostics.",
    )
    abstract = (
        "Introduction: Biomedical graph-learning results are difficult to compare when source versions, identifier rules, data splits, structural features, and negative-sampling contracts are incompletely documented. "
        "Methods: We developed BioGraphBench-PPI as an audit-first benchmark for protein-protein interaction (PPI) link prediction using STRING v12.0, BioGRID 5.0.260, and two cross-resource no-overlap ablations. For each PPI task, we generated 10 predefined seeds and three balanced negative-sampling regimes: random non-edges, degree-matched non-edges, and two-hop hard non-edges sampled from the training graph. We compared local heuristics, logistic regression, random forest, histogram gradient boosting (HGB), and an auditable node2vec-compatible random-walk baseline, reporting ranking, calibration, runtime, confidence intervals, paired Wilcoxon tests, and effect sizes. "
        "Results: Random negatives yielded high HGB AUPRC values (0.942-0.964), but degree-matched negatives reduced AUPRC by 0.049-0.259. Two-hop negatives changed the difficulty profile and model ranking, with HGB/RF AUPRC values of 0.774-0.857 and node2vec AUPRC values of 0.605-0.736. A retained functional-annotation pilot remained useful as a secondary stress test, but the manuscript's central empirical contribution is the PPI protocol. "
        "Discussion: BioGraphBench-PPI shows that PPI benchmark conclusions depend strongly on negative-sampling design and graph source. The benchmark therefore contributes an auditable evaluation contract rather than a claim that PPI prediction is solved."
    )
    set_text(doc.paragraphs[10], abstract)
    set_text(
        doc.paragraphs[11],
        "Keywords: protein-protein interactions, link prediction, negative sampling, node2vec, biomedical networks, benchmarking, reproducibility, calibration",
    )


def update_methods(doc: Document) -> None:
    replacements = {
        "We present BioGraphBench, an audit-first benchmark": (
            "We present BioGraphBench-PPI, a focused audit-first benchmark for reproducible learning on protein interaction networks. The central research question is whether open PPI resources can be transformed into link-prediction tasks whose evidence chain is reconstructible from source download through canonicalization, split generation, feature construction, model fitting, and statistical reporting."
        ),
        "The contribution is both methodological and empirical.": (
            "The contribution is both methodological and empirical. Methodologically, the benchmark makes data provenance, leakage control, negative sampling, and split contracts first-class experimental objects rather than post hoc documentation. Empirically, it establishes a demanding multi-seed non-GNN floor for PPI link prediction and shows that model ranking depends on the negative-sampling contract."
        ),
        "Five tasks were accepted": (
            "Five tasks were retained in the working benchmark, with four PPI link-prediction tasks forming the central contribution of this manuscript (Table 1). Their graph sizes and fixed positive partitions are summarized in Table 2. For each PPI graph, positive edges were shuffled with 10 predefined seeds (42-51) and partitioned into 80% training, 10% validation, and 10% testing subsets while preserving a spanning forest in the training graph. The OBNB task is retained as a secondary stress-test artifact, but the revised article's quantitative claims are centered on PPI link prediction."
        ),
        "All structural features were computed": (
            "All structural features were computed from Gtrain = (V, E+train) only. For each PPI partition and seed, three balanced negative-sampling contracts were produced. The random contract samples uniformly from unordered node pairs absent from the complete observed positive-edge set. The degree-matched contract approximately preserves the endpoint log-degree-bin distribution of the corresponding positive partition. The two-hop contract samples non-edges whose endpoints share at least one neighbor in the training-positive graph, creating a local hard-negative condition without using validation or test positives to define hardness."
        ),
        "Each individual heuristic served": (
            "Each individual heuristic served as an unsupervised score. The complete feature vector was supplied to logistic regression, random forest, and histogram gradient boosting (HGB) as supervised structural baselines. We additionally implemented an auditable node2vec-compatible random-walk baseline using unbiased walks (p = q = 1) and a logistic edge decoder trained on dot product, cosine similarity, L1 distance, and L2 distance features. Models were implemented with scikit-learn and PyTorch, and all model configurations and seeds are recorded in output manifests."
        ),
        "The software environment was pinned": (
            "The software environment was pinned to NumPy 1.24.3, pandas 2.0.3, SciPy 1.10.1, scikit-learn 1.3.2, PyTorch 2.4.1+CPU, and OBNB 0.1.0. Data checks, split validation, feature-policy tests, artifact-presence tests, raw per-seed outputs, statistical summaries, and regenerated figure/table scripts were retained with the benchmark. The central PPI results in the present manuscript come from the completed 10-seed Phase 1 protocol; the OBNB node-classification experiment is retained as a secondary challenge analysis."
        ),
    }
    for paragraph in doc.paragraphs:
        for prefix, text in replacements.items():
            if paragraph.text.startswith(prefix):
                set_text(paragraph, text)


def update_results_text(doc: Document) -> None:
    replacements = {
        "BioGraphBench retained five tasks": (
            "BioGraphBench retained five auditable tasks, with four PPI link-prediction tasks forming the central Phase 1 contribution (Table 1). The two main PPI graphs contained 18,767-20,376 nodes and 738,805-961,531 canonical edges. Removing cross-resource overlap reduced BioGRID to 538,003 edges and STRING to 314,539 edges, creating ablation tasks that test whether apparent performance is driven by shared database content. Across four PPI datasets, three negative-sampling contracts, and 10 seeds, the revised protocol produced 120 split manifests, 1,680 classical-baseline rows, and 240 node2vec-compatible rows (Table 2)."
        ),
        "Classical topology alone separated observed PPI edges": (
            "Classical topology and supervised structural baselines separated observed PPI edges from random non-edges effectively, but the revised multi-seed analysis shows that this result is conditional on the negative-sampling contract (Table 3). Under random negatives, HGB achieved mean test AUPRC values of 0.950 for STRING, 0.943 for BioGRID, 0.964 for STRING without BioGRID, and 0.942 for BioGRID without STRING. The seed-level distributions were tight, indicating reproducibility, but also confirming that random negatives define the easiest PPI contract."
        ),
        "Supervised combinations of the train-graph features": (
            "Replacing random negatives with degree-matched negatives materially reduced HGB performance in every PPI dataset. The absolute AUPRC drops were 0.069 for STRING, 0.172 for BioGRID, 0.049 for STRING without BioGRID, and 0.259 for BioGRID without STRING (Table 4). Two-hop hard negatives created a different profile: HGB AUPRC was 0.856 for STRING, 0.774 for BioGRID, 0.847 for STRING without BioGRID, and 0.849 for BioGRID without STRING. Thus, negative difficulty is not a single monotonic property; it depends on whether the contract controls endpoint degree, local closure, or both."
        ),
        "These high values do not demonstrate prospective biological completion.": (
            "These high values do not demonstrate prospective biological completion. They show that PPI link prediction is highly sensitive to how non-interacting pairs are constructed. Random negatives inflate apparent performance, degree-matched negatives reduce degree shortcuts, and two-hop negatives force models to distinguish observed interactions from structurally plausible local non-edges. The results therefore define an auditable baseline floor that future GNNs must exceed under all three contracts."
        ),
        "Ranking performance alone did not fully distinguish": (
            "Ranking performance alone did not fully distinguish the supervised models. The expanded model-family comparison shows that HGB and random forest are generally strongest, but not uniformly dominant. In BioGRID without STRING, node2vec-compatible embeddings exceeded HGB/RF under degree-matched negatives (AUPRC 0.725 versus 0.684-0.687), whereas HGB recovered the best value under two-hop negatives (0.849). In STRING without BioGRID, random forest slightly exceeded HGB under two-hop negatives (0.851 versus 0.847)."
        ),
        "The paired comparison in Figure 4A emphasizes": (
            "Calibration and runtime diagnostics add another layer to the benchmark. HGB generally showed lower ECE-10 than logistic regression and node2vec, while logistic regression remained the fastest baseline and random forest was the most expensive. These practical diagnostics matter because a benchmark baseline should be accurate, calibratable, and computationally interpretable, not only high ranking on AUPRC."
        ),
        "The results establish two different criteria": (
            "The results establish concrete requirements for future BioGraphBench-PPI submissions. A new GNN should be compared against HGB, random forest, and node2vec-compatible embeddings; should report mean, SD, confidence intervals, paired tests, and calibration; and should demonstrate that any gain persists across random, degree-matched, and two-hop negative contracts. The retained OBNB pilot continues to show that biological node-label prediction is a harder and qualitatively different problem, but it is not the central claim of the present PPI-focused article."
        ),
    }
    for paragraph in doc.paragraphs:
        for prefix, text in replacements.items():
            if paragraph.text.startswith(prefix):
                set_text(paragraph, text)

    for paragraph in doc.paragraphs:
        if "Figure 4A" in paragraph.text or "Figure 4B" in paragraph.text:
            paragraph.text.replace("Figure 4", "Figure 5")
        if paragraph.text.startswith("Validation-selected label thresholds"):
            set_text(
                paragraph,
                "Validation-selected label thresholds recovered non-zero test F1 but exposed a severe precision-recall trade-off (Figure 6). Logistic regression obtained the highest micro-F1 (0.0253) with precision 0.0131 and recall 0.4162. The MLP and GCN increased recall to 0.6878 and 0.6054, respectively, but precision remained near 0.011. This secondary result remains useful because it prevents BioGraphBench from being interpreted as only a PPI link-prediction exercise; however, the manuscript's central statistical evidence is the expanded PPI analysis.",
            )


def update_discussion_conclusion(doc: Document) -> None:
    replacements = {
        "BioGraphBench shows that the scientific value": (
            "BioGraphBench-PPI shows that the scientific value of a biomedical graph benchmark depends on the interpretability of its experimental contract. The revised PPI analysis demonstrates that high AUPRC under random negatives is reproducible but incomplete evidence: degree-matched and two-hop contracts can substantially change absolute performance and model ranking. The benchmark therefore evaluates not only models, but also the assumptions embedded in task construction."
        ),
        "The functional node-classification task prevents": (
            "The retained functional node-classification task prevents the benchmark from being interpreted as a single easy PPI exercise. Its sparse GOBP labels were not recovered effectively from degree bins or by the pilot GCN, confirming that different biological graph tasks expose different failure modes. In the present manuscript, however, this result should be read as a secondary stress test rather than as a fully developed benchmark family."
        ),
        "The present findings are consistent": (
            "The present findings are consistent with a broader literature in which graph structure can strongly support molecular interaction prediction, while evaluation design can change the apparent ranking of graph-learning methods. BioGraphBench-PPI extends that lesson by showing the size of the effect within open PPI resources: random, degree-matched, and two-hop negatives answer different questions, and a credible model claim must specify which question was asked."
        ),
        "Several limitations constrain": (
            "Several limitations constrain the current conclusions. First, the manuscript is now intentionally focused on PPI link prediction and should not be sold as a general network-bioinformatics benchmark. Second, the two-hop strategy is a local hard-negative proxy rather than a complete community-aware or temporally held-out evaluation. Third, the node2vec-compatible baseline is auditable but not exhaustively tuned. Fourth, competitive GNN link-prediction families such as GCN, GraphSAGE, GAT/GIN, and SEAL-style subgraph methods remain necessary for the next phase."
        ),
        "The next benchmark release": (
            "The next benchmark release should extend this Phase 1 protocol with competitive GNN link-prediction models, robustness tests under edge perturbation and STRING-threshold variation, model interpretability, and scalability measurements for memory, parameters, and inference time. Those extensions should preserve the same audit-first contract rather than replacing it with a model-centered leaderboard."
        ),
        "BioGraphBench converts open biomedical networks": (
            "BioGraphBench-PPI converts open protein-interaction resources into reconstructible experimental tasks and exposes a central methodological reality: PPI link-prediction conclusions depend strongly on the negative-sampling contract. Random negatives produced high and stable structural baselines, degree-matched negatives reduced degree shortcuts, and two-hop negatives changed both task difficulty and model ranking. The resulting benchmark is not a claim that PPI prediction is solved; it is a reproducible evidence framework for determining whether future graph models improve beyond structural shortcuts while remaining calibrated, auditable, and computationally interpretable."
        ),
        "BioGraphBench was developed from openly accessible": (
            "BioGraphBench-PPI was developed from openly accessible STRING, BioGRID, and OBNB resources. The working project snapshot contains source manifests, task definitions, 120 PPI split manifests, raw multi-seed outputs, statistical summaries, figures, tables, and validation records supporting this manuscript."
        ),
    }
    for paragraph in doc.paragraphs:
        for prefix, text in replacements.items():
            if paragraph.text.startswith(prefix):
                set_text(paragraph, text)


def update_figures(doc: Document) -> None:
    # Keep original Figure 1 methodological workflow to preserve the article architecture.
    replace_figure_before_caption(
        doc,
        "Figure 2 |",
        FIG_DIR / "fig1_hgb_auprc_seed_distributions.png",
        "Figure 2 | HGB performance distributions across 10 seeds. Violin densities and seed-level points summarize test AUPRC for complete PPI graphs and cross-database no-overlap ablations under random, degree-matched, and two-hop negative-sampling contracts.",
    )
    replace_figure_before_caption(
        doc,
        "Figure 3 |",
        FIG_DIR / "fig2_hgb_negative_regime_drops.png",
        "Figure 3 | Seed-paired performance change induced by harder negative sampling. Points show per-seed AUPRC differences between random negatives and the harder contracts; whiskers show 95% confidence intervals across 10 seeds.",
    )
    replace_figure_before_caption(
        doc,
        "Figure 4 |",
        FIG_DIR / "fig3_model_family_regime_profiles.png",
        "Figure 4 | Model-family profiles across negative-sampling regimes. Mean test AUPRC across 10 seeds is shown for logistic regression, random forest, HGB, and the node2vec-compatible random-walk baseline in each PPI task.",
    )
    add_figure_after_caption(
        doc,
        "Figure 4 |",
        FIG_DIR / "fig4_calibration_and_runtime.png",
        "Figure 5 | Calibration and computational-cost diagnostics. (A) ECE-10 distributions by model family and negative-sampling contract. (B) Training-time distributions by model family, with points showing dataset-level runs.",
    )
    fig6 = find_paragraph(doc, "Figure 5 | OBNB")
    set_text(
        fig6,
        "Figure 6 | OBNB BioGRID+GOBP node-classification trade-offs. (A) Test macro-AUROC and macro-AUPRC. (B) Test micro-F1, precision, and recall after per-label thresholds were selected on validation data.",
        "Caption",
    )


def update_tables(doc: Document, data: pd.DataFrame) -> None:
    test = data[data["split"] == "test"].copy()
    summary = test[test["model"].isin(MODEL_ORDER)].groupby(["dataset", "negative_strategy", "model"])["auprc"].mean().to_dict()

    # Table 3: replace old seed-42 heuristic table with multi-regime model table.
    t3 = doc.tables[2]
    table_set_size(t3, 13, 6)
    rows = [["Task", "Negatives", "LogReg", "RF", "HGB", "node2vec"]]
    for dataset in DATASET_ORDER:
        for neg in NEGATIVE_ORDER:
            rows.append(
                [
                    DATASET_LABELS[dataset],
                    NEGATIVE_LABELS[neg],
                    f"{summary[(dataset, neg, 'logistic_regression')]:.3f}",
                    f"{summary[(dataset, neg, 'random_forest')]:.3f}",
                    f"{summary[(dataset, neg, 'hist_gradient_boosting')]:.3f}",
                    f"{summary[(dataset, neg, 'node2vec_walk_logreg')]:.3f}",
                ]
            )
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t3.cell(r, c), value)

    # Table 4: replace old HGB single-regime calibration table with negative sensitivity.
    t4 = doc.tables[3]
    table_set_size(t4, 5, 6)
    rows = [["Task", "Random HGB", "Degree HGB", "Two-hop HGB", "Drop random-degree", "Drop random-two-hop"]]
    hgb = test[test["model"] == "hist_gradient_boosting"]
    for dataset in DATASET_ORDER:
        vals = {
            neg: float(hgb[(hgb["dataset"] == dataset) & (hgb["negative_strategy"] == neg)]["auprc"].mean())
            for neg in NEGATIVE_ORDER
        }
        rows.append(
            [
                DATASET_LABELS[dataset],
                f"{vals['random']:.3f}",
                f"{vals['degree_matched']:.3f}",
                f"{vals['two_hop']:.3f}",
                f"{vals['random'] - vals['degree_matched']:.3f}",
                f"{vals['random'] - vals['two_hop']:.3f}",
            ]
        )
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            fill_cell(t4.cell(r, c), value)

    captions = {
        "Table 1 | Accepted tasks": "Table 1 | Accepted tasks in the full BioGraphBench working release. The revised manuscript centers on the four PPI link-prediction tasks; OBNB is retained as a secondary stress-test artifact.",
        "Table 2 | Scale and fixed partitions": "Table 2 | Scale and fixed positive partitions of the accepted tasks. PPI counts are positive edges per seed; each partition is paired with an equal number of negatives under each PPI negative-sampling contract.",
        "Table 3 | Best classical heuristic": "Table 3 | Mean test AUPRC across 10 seeds by PPI dataset, negative-sampling contract, and model family.",
        "Table 4 | Calibration and recorded": "Table 4 | HGB sensitivity to negative-sampling contracts. Drops are seed-averaged AUPRC differences relative to random negatives.",
        "Table 5 | Test performance": "Table 5 | Test performance on OBNB BioGRID+GOBP multilabel node classification. Micro metrics use per-label thresholds selected on validation data. The constant control did not have a separate threshold-tuning record, so precision and recall are not reported.",
    }
    for paragraph in doc.paragraphs:
        for prefix, text in captions.items():
            if paragraph.text.startswith(prefix):
                set_text(paragraph, text, "Caption")


def main() -> int:
    source = INPUT_DOCX if INPUT_DOCX.exists() else FALLBACK_INPUT_DOCX
    data = read_results()
    doc = Document(source)
    update_front_and_abstract(doc)
    update_methods(doc)
    update_results_text(doc)
    update_discussion_conclusion(doc)
    update_figures(doc)
    update_tables(doc, data)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

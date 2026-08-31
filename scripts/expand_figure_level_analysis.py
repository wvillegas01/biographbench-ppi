"""Expand figure-level analysis in the integrated full-length manuscript."""

from __future__ import annotations

from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(r"C:\Users\wilop\Dropbox\MPDI\2026\BioGraphBench")
INPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised_full_length_tables_integrated.docx"
OUTPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised_full_length_tables_figures_integrated.docx"


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_after(paragraph, text: str, style: str | None = None):
    new_para = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_para._p)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def expand_figures_2_3(doc: Document) -> None:
    p = find_paragraph(doc, "The figures extend the table-level evidence")
    set_text(
        p,
        "Figure 2 extends Table 3 by showing that the PPI results are not driven by unstable seed behavior. In Figure 2A, the complete STRING and BioGRID graphs show extremely narrow HGB distributions across 10 seeds. STRING remains highest under random negatives (mean AUPRC 0.950), decreases under degree-matched negatives (0.881), and decreases further under two-hop negatives (0.856). BioGRID shows a sharper random-to-degree transition, falling from 0.943 to 0.770, while the two-hop condition remains similar to degree matching at 0.774. This panel therefore separates two issues: STRING is consistently easier than BioGRID under the harder contracts, and the seed variance is small relative to the protocol-induced shift.",
    )
    p2 = insert_after(
        p,
        "Figure 2B shows the no-overlap ablations and gives a stronger test of whether performance survives cross-resource decontamination. STRING without BioGRID preserves high performance under random negatives (0.964) and degree-matched negatives (0.916), but falls under two-hop negatives (0.847). BioGRID without STRING behaves differently: random negatives produce 0.942, degree-matched negatives drop sharply to 0.684, and two-hop negatives recover to 0.849. This pattern is important because it shows that harder negatives are not interchangeable. Degree matching and two-hop sampling stress different structural assumptions, especially after database overlap is removed.",
    )
    p3 = insert_after(
        p2,
        "Figure 3 makes the same result explicit as a paired contrast. In Figure 3A, random-minus-degree-matched differences are positive for all four datasets, confirming that degree matching systematically reduces HGB performance rather than introducing random noise. The largest reduction occurs in BioGRID without STRING (0.259), followed by BioGRID (0.172), STRING (0.069), and STRING without BioGRID (0.049). The ordering indicates that BioGRID-derived tasks are more vulnerable to endpoint-degree control than STRING-derived tasks.",
    )
    insert_after(
        p3,
        "Figure 3B compares random negatives with two-hop negatives. The reductions remain positive, but the order changes: BioGRID drops by 0.169, STRING without BioGRID by 0.117, STRING by 0.094, and BioGRID without STRING by 0.093. Compared with Figure 3A, this panel shows that two-hop negatives are especially disruptive for STRING-derived tasks and less disruptive than degree matching for BioGRID without STRING. Together, Figures 2 and 3 support the benchmark's central claim: BioGraphBench-PPI is reproducible across seeds, but sensitive enough to expose when the evaluation protocol is too easy.",
    )


def expand_figures_4_5(doc: Document) -> None:
    p = find_paragraph(doc, "Figure 4 moves beyond single-number performance")
    set_text(
        p,
        "Figure 4 analyzes model-family behavior one dataset at a time. In Figure 4A, STRING shows a smooth degradation from random to harder negatives: HGB and RF remain near the top, while node2vec and logistic regression fall more strongly under two-hop negatives. In Figure 4B, BioGRID shows a different profile: HGB and RF remain close under degree-matched and two-hop negatives, whereas node2vec remains lower, indicating that random-walk embeddings alone do not recover the same structural signal in the complete BioGRID graph.",
    )
    p2 = insert_after(
        p,
        "Figures 4C and 4D are the most informative panels because they use no-overlap ablations. In STRING without BioGRID (Figure 4C), HGB and RF are nearly indistinguishable under random and degree-matched negatives, and RF slightly exceeds HGB under two-hop negatives (0.851 versus 0.847). In BioGRID without STRING (Figure 4D), the ranking changes more dramatically: node2vec is strongest under degree-matched negatives (0.725), but HGB becomes strongest under two-hop negatives (0.849). These panel-level reversals show that model ranking is conditional on both data source and negative contract.",
    )
    p3 = find_paragraph(doc, "Figure 5 adds calibration and cost")
    set_text(
        p3,
        "Figure 5 adds practical diagnostics that cannot be inferred from AUPRC alone. In Figure 5A, logistic regression has the largest calibration error range, especially under harder negatives, whereas HGB remains tightly concentrated at low ECE-10 values. RF and node2vec occupy intermediate positions, with node2vec showing more calibration error than HGB despite being competitive in selected AUPRC settings. This panel supports reporting calibration as a primary companion to ranking metrics.",
    )
    insert_after(
        p3,
        "Figure 5B shows that runtime separates the baselines in a way that matters for benchmark adoption. Logistic regression is the fastest and most variable mainly because its cost scales simply with feature-table size. RF has the highest training-time distribution, reflecting the cost of ensemble fitting across large split files. HGB is less expensive than RF while retaining strong AUPRC and calibration, and node2vec has a relatively compact runtime band because the random-walk embedding stage dominates its cost profile. The practical conclusion is that HGB is a strong default baseline, but RF and node2vec remain necessary comparators because they represent different accuracy-cost trade-offs.",
    )


def expand_figure_6(doc: Document) -> None:
    p = find_paragraph(doc, "Figure 6 shows why the OBNB result")
    set_text(
        p,
        "Figure 6 should be read as a secondary stress-test figure rather than as part of the main PPI performance claim. In Figure 6A, macro-AUROC and macro-AUPRC remain low across the constant control, logistic regression, MLP, and GCN, showing that sparse GOBP annotation cannot be recovered effectively from simple degree-bin information or the pilot neural models. The contrast with Figures 2-4 is intentional: strong structural separability in PPI link prediction does not transfer automatically to multilabel functional annotation.",
    )
    insert_after(
        p,
        "Figure 6B shows the decision-level failure mode. Threshold tuning produces non-zero micro-F1, but the gains come with very low precision. Logistic regression obtains the highest micro-F1 (0.0253) by balancing low precision (0.0131) against moderate recall (0.4162), while MLP and GCN increase recall but do not solve the precision problem. This figure therefore supports a bounded interpretation of OBNB: it is useful evidence that BioGraphBench can expose harder biological tasks, but it should not dilute the manuscript's central PPI-focused contribution.",
    )


def main() -> int:
    doc = Document(INPUT_DOCX)
    expand_figures_2_3(doc)
    expand_figures_4_5(doc)
    expand_figure_6(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

"""Integrate tables into the manuscript body and expand analytical results prose."""

from __future__ import annotations

from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(r"C:\Users\wilop\Dropbox\MPDI\2026\BioGraphBench")
INPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised_full_length.docx"
OUTPUT_DOCX = PROJECT_ROOT / "BioGraphBench_Frontiers_working_manuscript_PPI_revised_full_length_tables_integrated.docx"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_text(paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    if style:
        paragraph.style = style
    paragraph.add_run(text)


def insert_after(paragraph, text: str, style: str | None = None):
    new_para = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_para._p)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def move_table_with_caption_after(doc: Document, caption_prefix: str, table_index: int, anchor):
    caption = find_paragraph(doc, caption_prefix)
    table = doc.tables[table_index]
    anchor._p.addnext(table._element)
    anchor._p.addnext(caption._p)
    return table


def remove_tables_heading_if_empty(doc: Document) -> None:
    try:
        heading = find_paragraph(doc, "Tables")
    except ValueError:
        return
    delete_paragraph(heading)


def expand_accepted_tasks(doc: Document) -> None:
    p = find_paragraph(doc, "BioGraphBench retained five auditable tasks")
    set_text(
        p,
        "The accepted task set should be read as a quality-control result, not only as a list of available datasets. BioGraphBench retained five auditable tasks because their source files, transformations, split definitions, feature policies, and outputs could be reconstructed from local manifests. Table 1 separates the central PPI contribution from the retained OBNB stress test: four tasks evaluate undirected PPI link prediction, whereas OBNB remains a harder secondary node-labeling task that protects the benchmark from being interpreted as a single easy graph problem.",
    )
    p2 = insert_after(
        p,
        "This distinction matters for the article's scope. The revised manuscript is no longer framed as a broad benchmark claim across all network-bioinformatics tasks. Instead, it uses the full working release to justify dataset auditability, then focuses its statistical evidence on PPI link prediction. Table 1 therefore performs two roles: it documents the larger benchmark scaffold and defines the empirical boundary of the present paper.",
    )
    table1 = move_table_with_caption_after(doc, "Table 1 |", 0, p2)
    after_t1 = insert_after(
        p2,
        "The scale summary in Table 2 shows why the PPI study is a meaningful stress test even before model comparison. The complete STRING and BioGRID graphs are not small illustrative examples; they contain hundreds of thousands of canonical interactions. The no-overlap ablations are also large enough to support evaluation, but they remove shared cross-resource pairs that could otherwise make performance appear more general than it is. This makes the ablations essential rather than decorative: they test whether a model is learning robust structural regularities or benefiting from database overlap.",
    )
    # Place paragraph after the moved table element by anchoring on the caption again.
    table1._element.addnext(after_t1._p)
    table2 = move_table_with_caption_after(doc, "Table 2 |", 1, after_t1)
    after_t2 = insert_after(
        after_t1,
        "Across four PPI datasets, three negative-sampling contracts, and 10 seeds, the protocol generated 120 split manifests. This is the first major strengthening over the earlier pilot: the benchmark is no longer only deterministic, it now quantifies sampling stability. The 1,680 classical-baseline rows and 240 node2vec-compatible rows are therefore not bookkeeping details; they are the empirical basis for confidence intervals, paired comparisons, and model-ranking claims.",
    )
    table2._element.addnext(after_t2._p)


def expand_structural_results(doc: Document) -> None:
    p = find_paragraph(doc, "Classical topology and supervised structural baselines")
    set_text(
        p,
        "Table 3 is the central performance table because it replaces a single-protocol result with a contract-aware comparison. Under random negatives, all supervised structural models perform strongly, especially HGB and random forest. HGB reaches mean test AUPRC values of 0.950 for STRING, 0.943 for BioGRID, 0.964 for STRING without BioGRID, and 0.942 for BioGRID without STRING. If this were the only result, the benchmark would appear nearly solved by local topology.",
    )
    p2 = insert_after(
        p,
        "The same table shows why that interpretation would be misleading. Degree-matched negatives reduce the shortcut provided by endpoint degree, while two-hop negatives force discrimination among locally plausible pairs. The model rankings are no longer invariant: HGB and random forest remain strong in most contexts, but node2vec becomes competitive in the BioGRID-without-STRING degree-matched condition, and random forest slightly exceeds HGB for STRING-without-BioGRID under two-hop negatives. This is a methodological result, not just a model result: benchmark conclusions change when the negative contract changes.",
    )
    table3 = move_table_with_caption_after(doc, "Table 3 |", 2, p2)

    p_neg = find_paragraph(doc, "Replacing random negatives with degree-matched negatives")
    set_text(
        p_neg,
        "Table 4 quantifies the cost of replacing random negatives with harder contracts. Degree-matched negatives reduce HGB AUPRC by 0.069 in STRING, 0.172 in BioGRID, 0.049 in STRING without BioGRID, and 0.259 in BioGRID without STRING. The largest drop occurs in BioGRID without STRING, which is precisely the setting where overlap has been removed and degree matching most strongly challenges the structural baseline.",
    )
    p_neg2 = insert_after(
        p_neg,
        "The two-hop contract tells a complementary story. Its performance loss relative to random negatives is substantial in STRING and STRING without BioGRID, but it does not behave identically to degree matching. In BioGRID without STRING, two-hop HGB AUPRC rises to 0.849, higher than the degree-matched value of 0.684. This reversal is important: hard negatives are not interchangeable. Degree matching controls endpoint popularity, whereas two-hop sampling emphasizes local closure. A defensible benchmark should report both because they test different failure modes.",
    )
    table4 = move_table_with_caption_after(doc, "Table 4 |", 3, p_neg2)

    # Keep subsequent figure discussion after the integrated tables.
    p_high = find_paragraph(doc, "These high values do not demonstrate prospective biological completion")
    set_text(
        p_high,
        "The figures extend the table-level evidence by showing seed stability and paired contrasts. Figure 2 shows that HGB results are very stable across 10 seeds, so the main concern is not random instability. Instead, Figure 3 shows that the negative-sampling contract systematically shifts performance. In other words, the benchmark is reproducible and sensitive: it is stable enough to audit, but sensitive enough to expose when an evaluation protocol is too easy.",
    )


def expand_calibration_and_obnb(doc: Document) -> None:
    p = find_paragraph(doc, "Ranking performance alone did not fully distinguish")
    set_text(
        p,
        "Figure 4 moves beyond single-number performance by showing how model families respond across negative contracts. The most useful pattern is not that one model always wins; rather, the figure shows where rankings are stable and where they change. HGB and random forest form the strongest practical non-GNN baselines in most settings, but node2vec provides an important embedding-based comparator and prevents the benchmark from being reduced to handcrafted pair features.",
    )
    p2 = find_paragraph(doc, "Calibration and runtime diagnostics add another layer")
    set_text(
        p2,
        "Figure 5 adds calibration and cost, which are necessary for benchmark interpretation. HGB generally combines strong ranking with low ECE-10, whereas logistic regression is faster but less competitive under harder contracts. Random forest is often competitive in AUPRC but more expensive, and node2vec occupies a different cost-performance profile because it requires random-walk embedding before decoding. These diagnostics support the benchmark's practical recommendation: future GNNs should be compared against baselines that are not only accurate, but also calibrated and computationally transparent.",
    )

    p3 = find_paragraph(doc, "The multilabel OBNB BioGRID+GOBP task")
    set_text(
        p3,
        "The multilabel OBNB BioGRID+GOBP task remains in the manuscript as a secondary stress test rather than as part of the main Phase 1 PPI claim (Table 5). Its performance regime is deliberately different. A constant-feature logistic regression control gave the chance-level macro-AUROC of 0.500 and macro-AUPRC 0.0107. One-hot log-degree features raised logistic-regression macro-AUROC to 0.531 and macro-AUPRC to 0.0144, while the pilot MLP and GCN did not produce a decisive improvement. This contrast helps show that strong PPI link-prediction baselines do not imply that all biomedical graph tasks are easy.",
    )
    table5 = move_table_with_caption_after(doc, "Table 5 |", 4, p3)
    p4 = find_paragraph(doc, "Validation-selected label thresholds")
    set_text(
        p4,
        "Figure 6 shows why the OBNB result should be interpreted cautiously but retained. Validation-selected label thresholds recovered non-zero test F1, yet exposed a severe precision-recall trade-off. Logistic regression obtained the highest micro-F1 (0.0253) with precision 0.0131 and recall 0.4162. The MLP and GCN increased recall to 0.6878 and 0.6054, respectively, but precision remained near 0.011. This secondary analysis strengthens the discussion by showing that BioGraphBench can expose task difficulty beyond PPI, while the present paper remains appropriately focused on PPI link prediction.",
    )
    table5._element.addnext(p4._p)


def update_discussion(doc: Document) -> None:
    p = find_paragraph(doc, "BioGraphBench-PPI shows that the scientific value")
    set_text(
        p,
        "BioGraphBench-PPI shows that the scientific value of a biomedical graph benchmark depends on the interpretability of its experimental contract. The integrated tables and figures support a single coherent conclusion: high PPI link-prediction performance under random negatives is real and reproducible, but incomplete. Degree-matched and two-hop contracts change both the absolute level of performance and the relative ranking of model families. This means the benchmark is not simply measuring model capacity; it is also measuring the assumptions embedded in task construction.",
    )


def main() -> int:
    doc = Document(INPUT_DOCX)
    expand_accepted_tasks(doc)
    expand_structural_results(doc)
    expand_calibration_and_obnb(doc)
    update_discussion(doc)
    remove_tables_heading_if_empty(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

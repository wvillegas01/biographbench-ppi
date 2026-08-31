"""Apply reviewer-driven textual corrections to the submitted manuscript in red.

The script edits the submitted manuscript in place after a dated backup has
already been created. Only modified/replaced text is colored red so changes are
easy to locate in Word.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor


DOCX_PATH = Path(r"C:\Users\wilop\Dropbox\MPDI\2026\BioGraphBench\1966469_Manuscript.DOCX")
LOG_PATH = Path(
    r"C:\Users\wilop\Dropbox\MPDI\2026\BioGraphBench\Auditoria\Articulo\reviewer_text_corrections_red_20260831.md"
)

RED = RGBColor(192, 0, 0)


def red_replace(paragraph, text: str, style: str | None = None) -> None:
    paragraph.clear()
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    run.font.color.rgb = RED


def insert_red_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    run = new_para.add_run(text)
    run.font.color.rgb = RED
    return new_para


def find_paragraph(doc: Document, needle: str):
    for idx, paragraph in enumerate(doc.paragraphs):
        if needle.lower() in paragraph.text.lower():
            return idx, paragraph
    raise ValueError(f"Could not find paragraph containing: {needle}")


def main() -> int:
    doc = Document(DOCX_PATH)
    changes: list[str] = []

    # Reviewer 2, comment 4: explain notation.
    _, notation_explanation = find_paragraph(doc, "where Gₖ is the canonical graph")
    red_replace(
        notation_explanation,
        "Here, k indexes a dataset-task instance. Gₖ denotes the canonical graph for task k; Tₖ is the prediction task; Sₖ contains the predefined partitions and sampled-nonedge contracts; Φₖ specifies the feature-construction policy; and Mₖ defines the metric set. Within Gₖ, Vₖ is the node set, Eₖ the edge set, Xₖ the input representation, and Yₖ the labels or prediction targets. An experiment is accepted only when its source, transformation rules, splits, sampled-nonedge contract, features, model configuration, seed, and outputs have machine-readable artifacts. Figure 1 summarizes the BioGraphBench-PPI evidence chain.",
    )
    changes.append("Expanded notation explanation after the benchmark-object expression.")

    # Reviewer 1, comment 2: clarify 738-pair gap.
    _, overlap_para = find_paragraph(doc, "The two resources shared 423,528 canonical pairs")
    red_replace(
        overlap_para,
        "To quantify direct reuse between sources, interaction pairs were aligned in Entrez Gene space for cross-resource comparison. Overlap statistics were calculated on these aligned Entrez pair sets. The two resources shared 423,528 canonical Entrez pairs, corresponding to 55.27% of the mapped STRING pair set and 44.05% of the mapped BioGRID pair set, with a Jaccard index of 0.3247. Reciprocal ablations then removed source-specific edges that mapped to an overlapping Entrez pair. For BioGRID, this removed 423,528 Entrez edges. For STRING, this removed 424,266 original STRING protein-pair edges. The 738-edge difference arises because STRING protein identifiers may map to Entrez identifiers through multiple explicit aliases; therefore, the Entrez-space overlap count and the STRING-space ablation count are related but not identical. For sources a and b,",
    )
    changes.append("Clarified the 423,528 versus 424,266 overlap/ablation count difference.")

    # Reviewer 2, comment 2: avoid undefined Phase 1.
    for needle, replacement in [
        (
            "completed 10-seed Phase 1 protocol",
            "The software environment was pinned to NumPy 1.24.3, pandas 2.0.3, SciPy 1.10.1, scikit-learn 1.3.2, PyTorch 2.4.1+CPU, and OBNB 0.1.0. Data checks, split validation, feature-policy tests, artifact-presence tests, raw per-seed outputs, statistical summaries, and regenerated figure/table scripts were retained with the benchmark. The central PPI results in the present benchmark release come from the completed 10-seed protocol; the OBNB node-classification experiment is retained as a secondary challenge analysis.",
        ),
        (
            "main Phase 1 PPI claim",
            "The multilabel OBNB BioGRID+GOBP task remains in the manuscript as a secondary stress test rather than as part of the main PPI link-prediction claim (Table 5). Its performance regime is deliberately different. A constant-feature logistic regression control gave the chance-level macro-AUROC of 0.500 and macro-AUPRC 0.0107. One-hot log-degree features raised logistic-regression macro-AUROC to 0.531 and macro-AUPRC to 0.0144, while the pilot MLP and GCN did not produce a decisive improvement. This contrast helps show that strong PPI link-prediction baselines do not imply that all biomedical graph tasks are easy.",
        ),
        (
            "extend this Phase 1 protocol",
            "The next benchmark release should extend the present PPI protocol with competitive GNN link-prediction models, robustness tests under edge perturbation and STRING-threshold variation, model interpretability, and scalability measurements for memory, parameters, and inference time. Those extensions should preserve the same audit-first contract rather than replacing it with a model-centered leaderboard.",
        ),
    ]:
        _, para = find_paragraph(doc, needle)
        red_replace(para, replacement)
        changes.append(f"Replaced undefined/revision-history wording around: {needle}")

    # Reviewer 2, comment 1: remove reader-inappropriate phrasing.
    _, limitations = find_paragraph(doc, "should not be sold as")
    red_replace(
        limitations,
        "Several limitations constrain the current conclusions. First, the present study is intentionally scoped to PPI link prediction; broader network-bioinformatics tasks require additional task-specific validation. Second, the two-hop strategy is a local hard-negative proxy rather than a complete community-aware or temporally held-out evaluation. Third, the node2vec-compatible baseline is auditable but not exhaustively tuned. Fourth, competitive GNN link-prediction families such as GCN, GraphSAGE, GAT/GIN, and SEAL-style subgraph methods remain necessary for a model-centered extension of the benchmark.",
    )
    changes.append("Removed 'now' and 'sold as' from limitations.")

    _, dilute = find_paragraph(doc, "should not dilute")
    red_replace(
        dilute,
        "Figure 6B shows the decision-level failure mode. Threshold tuning produces non-zero micro-F1, but the gains come with very low precision. Logistic regression obtains the highest micro-F1 (0.0253) by balancing low precision (0.0131) against moderate recall (0.4162), while MLP and GCN increase recall but do not solve the precision problem. This figure therefore supports a bounded interpretation of OBNB: it is useful evidence that BioGraphBench can expose harder biological tasks, but it is interpreted as a secondary stress test rather than as part of the central PPI link-prediction claim.",
    )
    changes.append("Removed reader-inappropriate 'dilute the manuscript' phrasing.")

    # Reviewer 2, comment 6: moderate unsupported future-graph-model claim.
    _, req_heading = find_paragraph(doc, "Empirical requirements for future graph models")
    red_replace(req_heading, "Empirical requirements for future model submissions", style=req_heading.style.name)
    changes.append("Renamed the future-graph-model heading to avoid overclaiming.")

    _, req_para = find_paragraph(doc, "A new GNN should be compared")
    red_replace(
        req_para,
        "The results establish concrete requirements for future BioGraphBench-PPI submissions. Any new link-prediction model, including GNN-based models, should be compared against HGB, random forest, and node2vec-compatible embeddings; should report mean, SD, confidence intervals, paired seed-wise comparisons, and calibration; and should demonstrate that any gain persists across random, degree-matched, and two-hop sampled-nonedge contracts. The retained OBNB pilot continues to show that biological node-label prediction is a harder and qualitatively different problem, but it is not the central empirical claim of the present PPI-focused article.",
    )
    changes.append("Moderated GNN language while preserving benchmark requirements.")

    _, conclusion = find_paragraph(doc, "future graph models improve beyond structural shortcuts")
    red_replace(
        conclusion,
        "BioGraphBench-PPI converts open protein-interaction resources into reconstructible experimental tasks and exposes a central methodological reality: PPI link-prediction conclusions depend strongly on the sampled-nonedge contract. Random sampled nonedges produced high and stable structural baselines, degree-matched sampled nonedges reduced degree shortcuts, and two-hop sampled nonedges changed both task difficulty and model ranking. The resulting benchmark is not a claim that PPI prediction is solved; it is a reproducible evidence framework for testing whether future link-prediction models improve beyond structural and embedding baselines while remaining calibrated, auditable, and computationally interpretable.",
    )
    changes.append("Moderated the conclusion claim about future graph models.")

    # Reviewer 2, comment 5: data/code availability.
    _, availability = find_paragraph(doc, "available from the corresponding author")
    red_replace(
        availability,
        "The source datasets used in this study are publicly available through STRING, BioGRID, and OBNB. The BioGraphBench-PPI code, acquisition scripts, canonicalization scripts, split-generation scripts, split manifests, multi-seed result tables, statistical summaries, and figure-generation scripts will be deposited in a public repository before resubmission: [repository URL / DOI to be inserted before upload]. Raw STRING, BioGRID, and OBNB files should be obtained from their original providers according to their respective terms; the repository will provide checksums, manifests, and executable scripts required to reproduce the analyses.",
    )
    changes.append("Updated data/code availability statement with a public-repository placeholder.")

    # Reviewer 2, comment 3: complete node2vec title.
    _, node2vec = find_paragraph(doc, "Grover, A., and Leskovec")
    red_replace(
        node2vec,
        "Grover, A., and Leskovec, J. (2016). node2vec: Scalable feature learning for networks. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 855-864. doi: 10.1145/2939672.2939754.",
        style=node2vec.style.name,
    )
    changes.append("Completed the node2vec reference title.")

    # Make statistical test language explicit because reviewers asked for support.
    _, stats_para = find_paragraph(doc, "Student t distribution with 9 degrees")
    red_replace(
        stats_para,
        "For each PPI dataset-contract-model combination, performance metrics were summarized as the mean and sample standard deviation across 10 seeds. Ninety-five percent confidence intervals for the corresponding means were calculated using the Student t distribution with 9 degrees of freedom. Effects of the sampled-nonedge contract were evaluated as seed-paired differences between random and degree-matched conditions and between random and two-hop conditions, using the same seed in each comparison. Paired Wilcoxon signed-rank tests and paired standardized effect sizes were computed to support contract-level comparisons.",
    )
    changes.append("Made Wilcoxon/effect-size reporting explicit in Methods.")

    doc.save(DOCX_PATH)

    LOG_PATH.write_text(
        "# Reviewer Text Corrections Applied in Red\n\n"
        f"Edited file: `{DOCX_PATH}`\n\n"
        f"Backup file: `C:\\Users\\wilop\\Dropbox\\MPDI\\2026\\BioGraphBench\\1966469_Manuscript_backup_20260831.docx`\n\n"
        "## Changes\n\n"
        + "\n".join(f"- {item}" for item in changes)
        + "\n\n## Note\n\n"
        "No GNN link-prediction performance values were added because no such experiment has been run yet in the available result files. The manuscript language was moderated so current claims remain supported by the existing classical, node2vec-compatible, multi-seed, and sampled-nonedge experiments.\n",
        encoding="utf-8",
    )
    print(DOCX_PATH)
    print(LOG_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

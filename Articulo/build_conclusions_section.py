from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
OUT_MD = ARTICLE_DIR / "seccion_conclusiones.md"
OUT_DOCX = ARTICLE_DIR / "seccion_conclusiones_biographbench.docx"


CONCLUSIONS = """# 6. CONCLUSIONES

Este trabajo presentó BioGraphBench como una propuesta de benchmark reproducible para aprendizaje sobre grafos biomédicos, construido desde una premisa deliberada: la confianza experimental debe preceder a la complejidad algorítmica. La contribución principal no es un nuevo modelo ni una arquitectura específica, sino un marco auditable para transformar recursos biológicos abiertos en tareas evaluables, documentadas y comparables. En lugar de asumir que todo dataset descargable puede convertirse automáticamente en benchmark, BioGraphBench exige trazabilidad de origen, reglas explícitas de filtrado, definición formal de tareas, particiones reproducibles, features bajo restricciones de entrenamiento y métricas alineadas con el uso biomédico.

El estudio permitió consolidar un MVP con tareas de predicción de enlaces PPI y clasificación funcional multilabel, integrando recursos ampliamente utilizados como STRING, BioGRID y OBNB. Este alcance inicial es acotado, pero suficiente para demostrar la utilidad del enfoque. La estructura propuesta permite reconstruir cada resultado desde los datos crudos hasta los reportes finales, reduciendo ambigüedades frecuentes en estudios donde los pasos intermedios quedan implícitos. En ese sentido, BioGraphBench funciona como un contrato experimental: define qué se mide, con qué datos, bajo qué restricciones y contra qué baselines.

Una conclusión central es que la reproducibilidad no debe entenderse como una fase posterior al modelado, sino como una propiedad de diseño del benchmark. Cuando tareas, splits y features se formalizan antes de entrenar modelos, la evaluación deja de depender de decisiones ad hoc y se vuelve inspeccionable por otros investigadores. Esto es especialmente relevante en biomedicina, donde los datasets combinan evidencia incompleta, sesgos de curación, identificadores heterogéneos y relaciones solapadas entre fuentes. En este contexto, un benchmark útil no solo debe producir rankings, sino también explicar por qué son confiables.

BioGraphBench también aporta una base para evitar afirmaciones excesivas. El artículo no plantea que el MVP resuelva definitivamente la evaluación de GNNs biomédicas, sino que establece una infraestructura inicial para hacer esa evaluación más rigurosa. Su valor está en convertir preguntas dispersas sobre datasets, leakage, features, calibración y baselines en un pipeline verificable. Esta orientación permite que futuras comparaciones de modelos sean menos dependientes del entusiasmo arquitectónico y más cercanas a una evaluación científica acumulativa.

Como trabajos futuros, se propone ampliar el benchmark en cuatro direcciones. Primero, incorporar múltiples semillas, intervalos de confianza y pruebas estadísticas pareadas. Segundo, añadir nuevas familias de GNNs y modelos híbridos bajo el mismo protocolo. Tercero, extender las tareas hacia otros recursos biomédicos abiertos, manteniendo controles de licencia, descarga y solapamiento. Finalmente, publicar el pipeline como paquete instalable con documentación, versionado de datasets y scripts de reproducción completa."""


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("2E74B5")

    for line in CONCLUSIONS.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def main() -> None:
    OUT_MD.write_text(CONCLUSIONS, encoding="utf-8")
    build_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()

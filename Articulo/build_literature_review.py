from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
OUT_MD = ARTICLE_DIR / "seccion_revision_literatura.md"
OUT_DOCX = ARTICLE_DIR / "seccion_revision_literatura_biographbench.docx"
OUT_CSV = ARTICLE_DIR / "literature_review_references_20.csv"


REFERENCES = [
    [1, "Hu et al.", 2020, "Open Graph Benchmark: Datasets for Machine Learning on Graphs", "10.48550/arXiv.2005.00687", "https://arxiv.org/abs/2005.00687", "Si"],
    [2, "Dwivedi et al.", 2023, "Benchmarking Graph Neural Networks", "10.5555/3648699.3648742", "https://jmlr.org/papers/v24/22-0567.html", "Si"],
    [3, "Huang et al.", 2022, "Artificial intelligence foundation for therapeutic science", "10.1038/s41589-022-01131-2", "https://pmc.ncbi.nlm.nih.gov/articles/PMC9529840/", "Si"],
    [4, "Wu et al.", 2018, "MoleculeNet: a benchmark for molecular machine learning", "10.1039/C7SC02664A", "https://pubs.rsc.org/en/content/articlelanding/2018/sc/c7sc02664a", "Si"],
    [5, "Breit et al.", 2020, "OpenBioLink: a benchmarking framework for large-scale biomedical link prediction", "10.1093/bioinformatics/btaa274", "https://academic.oup.com/bioinformatics/article/36/13/4097/5825726", "Si"],
    [6, "Liu and Krishnan", 2024, "Open Biomedical Network Benchmark: A Python Toolkit for Benchmarking Datasets with Biomedical Networks", "10.1101/2023.01.10.523485", "https://proceedings.mlr.press/v240/liu24a.html", "Si"],
    [7, "Yue et al.", 2020, "Graph embedding on biomedical networks: methods, applications and evaluations", "10.1093/bioinformatics/btz718", "https://pmc.ncbi.nlm.nih.gov/articles/PMC7703771/", "Si"],
    [8, "Szklarczyk et al.", 2023, "The STRING database in 2023", "10.1093/nar/gkac1000", "https://academic.oup.com/nar/article/51/D1/D638/6825348", "Si"],
    [9, "Oughtred et al.", 2021, "The BioGRID database", "10.1002/pro.3978", "https://doi.org/10.1002/pro.3978", "Si"],
    [10, "Himmelstein et al.", 2017, "Systematic integration of biomedical knowledge prioritizes drugs for repurposing", "10.7554/eLife.26726", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5640425/", "Si"],
    [11, "Jha et al.", 2022, "Prediction of protein-protein interaction using graph neural networks", "10.1038/s41598-022-12201-9", "https://www.nature.com/articles/s41598-022-12201-9", "Si"],
    [12, "Mohamed et al.", 2022, "Graph Neural Network for Protein-Protein Interaction Prediction", "10.3390/molecules27186135", "https://www.mdpi.com/1420-3049/27/18/6135", "Si"],
    [13, "Lv et al.", 2024, "GNNGL-PPI: multi-category prediction of protein-protein interactions", "10.1186/s12864-024-10299-x", "https://bmcbioinformatics.biomedcentral.com/articles/10.1186/s12864-024-10299-x", "Si"],
    [14, "Huang et al.", 2020, "SkipGNN: predicting molecular interactions with skip-graph networks", "10.1038/s41598-020-77766-9", "https://www.nature.com/articles/s41598-020-77766-9", "Si"],
    [15, "Wu et al.", 2022, "BridgeDPI: a novel Graph Neural Network for predicting drug-protein interactions", "10.1093/bioinformatics/btac155", "https://pubmed.ncbi.nlm.nih.gov/35274672/", "Si"],
    [16, "Sun et al.", 2020, "Graph convolutional networks for computational drug development and discovery", "10.1093/bib/bbz042", "https://academic.oup.com/bib/article/21/3/919/5498046", "Parcial"],
    [17, "Feng et al.", 2024, "A review on graph neural networks for predicting synergistic drug combinations", "10.1007/s10462-023-10669-z", "https://link.springer.com/article/10.1007/s10462-023-10669-z", "Parcial"],
    [18, "Wan et al.", 2024, "Knowledge mapping of graph neural networks for drug discovery", "10.3389/fphar.2024.1393415", "https://pmc.ncbi.nlm.nih.gov/articles/PMC11116974/", "Si"],
    [19, "Grover and Leskovec", 2016, "node2vec: Scalable Feature Learning for Networks", "10.1145/2939672.2939754", "https://dl.acm.org/doi/10.1145/2939672.2939754", "Si"],
    [20, "Shchur et al.", 2018, "Pitfalls of Graph Neural Network Evaluation", "10.48550/arXiv.1811.05868", "https://arxiv.org/abs/1811.05868", "Si"],
]


APA_REFERENCES = [
    "Breit, A., Ott, S., Agibetov, A., & Samwald, M. (2020). OpenBioLink: A benchmarking framework for large-scale biomedical link prediction. Bioinformatics, 36(13), 4097-4098. https://doi.org/10.1093/bioinformatics/btaa274",
    "Dwivedi, V. P., Joshi, C. K., Luu, A. T., Laurent, T., Bengio, Y., & Bresson, X. (2023). Benchmarking graph neural networks. Journal of Machine Learning Research, 24, 1-48. https://doi.org/10.5555/3648699.3648742",
    "Feng, Y., Zhang, Y., & Wang, X. (2024). A review on graph neural networks for predicting synergistic drug combinations. Artificial Intelligence Review, 57. https://doi.org/10.1007/s10462-023-10669-z",
    "Grover, A., & Leskovec, J. (2016). node2vec: Scalable feature learning for networks. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 855-864. https://doi.org/10.1145/2939672.2939754",
    "Himmelstein, D. S., Lizee, A., Hessler, C., Brueggeman, L., Chen, S. L., Hadley, D., Green, A., Khankhanian, P., & Baranzini, S. E. (2017). Systematic integration of biomedical knowledge prioritizes drugs for repurposing. eLife, 6, e26726. https://doi.org/10.7554/eLife.26726",
    "Hu, W., Fey, M., Zitnik, M., Dong, Y., Ren, H., Liu, B., Catasta, M., & Leskovec, J. (2020). Open Graph Benchmark: Datasets for machine learning on graphs. arXiv. https://doi.org/10.48550/arXiv.2005.00687",
    "Huang, K., Fu, T., Gao, W., Zhao, Y., Roohani, Y., Leskovec, J., Coley, C. W., Xiao, C., Sun, J., & Zitnik, M. (2022). Artificial intelligence foundation for therapeutic science. Nature Chemical Biology, 18(10), 1033-1036. https://doi.org/10.1038/s41589-022-01131-2",
    "Huang, K., Xiao, C., Glass, L. M., & Sun, J. (2020). SkipGNN: Predicting molecular interactions with skip-graph networks. Scientific Reports, 10, 21092. https://doi.org/10.1038/s41598-020-77766-9",
    "Jha, K., Saha, S., & Singh, H. (2022). Prediction of protein-protein interaction using graph neural networks. Scientific Reports, 12, 8360. https://doi.org/10.1038/s41598-022-12201-9",
    "Liu, R., & Krishnan, A. (2024). Open Biomedical Network Benchmark: A Python toolkit for benchmarking datasets with biomedical networks. Proceedings of Machine Learning Research, 240, 23-59. https://doi.org/10.1101/2023.01.10.523485",
    "Lv, G., Hu, Z., Bi, Y., & Zhang, S. (2024). GNNGL-PPI: Multi-category prediction of protein-protein interactions using graph neural networks based on global graphs and local subgraphs. BMC Genomics, 25, 431. https://doi.org/10.1186/s12864-024-10299-x",
    "Mohamed, S. K., Nounu, A., & Nováček, V. (2022). Graph neural network for protein-protein interaction prediction. Molecules, 27(18), 6135. https://doi.org/10.3390/molecules27186135",
    "Oughtred, R., Rust, J., Chang, C., Breitkreutz, B. J., Stark, C., Willems, A., Boucher, L., Leung, G., Kolas, N., Zhang, F., Dolma, S., Coulombe-Huntington, J., Chatr-Aryamontri, A., Dolinski, K., & Tyers, M. (2021). The BioGRID database: A comprehensive biomedical resource of curated protein, genetic, and chemical interactions. Protein Science, 30(1), 187-200. https://doi.org/10.1002/pro.3978",
    "Shchur, O., Mumme, M., Bojchevski, A., & Günnemann, S. (2018). Pitfalls of graph neural network evaluation. arXiv. https://doi.org/10.48550/arXiv.1811.05868",
    "Sun, M., Zhao, S., Gilvary, C., Elemento, O., Zhou, J., & Wang, F. (2020). Graph convolutional networks for computational drug development and discovery. Briefings in Bioinformatics, 21(3), 919-935. https://doi.org/10.1093/bib/bbz042",
    "Szklarczyk, D., Kirsch, R., Koutrouli, M., Nastou, K., Mehryary, F., Hachilif, R., Gable, A. L., Fang, T., Doncheva, N. T., Pyysalo, S., Bork, P., Jensen, L. J., & von Mering, C. (2023). The STRING database in 2023: Protein-protein association networks and functional enrichment analyses for any sequenced genome of interest. Nucleic Acids Research, 51(D1), D638-D646. https://doi.org/10.1093/nar/gkac1000",
    "Wan, F., Zhang, Y., & Wu, Z. (2024). Knowledge mapping of graph neural networks for drug discovery: A bibliometric and visualized analysis. Frontiers in Pharmacology, 15, 1393415. https://doi.org/10.3389/fphar.2024.1393415",
    "Wu, Y., Gao, M., Zeng, M., Chen, F., Li, M., & Zhang, J. (2022). BridgeDPI: A novel graph neural network for predicting drug-protein interactions. Bioinformatics, 38(9), 2571-2578. https://doi.org/10.1093/bioinformatics/btac155",
    "Wu, Z., Ramsundar, B., Feinberg, E. N., Gomes, J., Geniesse, C., Pappu, A. S., Leswing, K., & Pande, V. (2018). MoleculeNet: A benchmark for molecular machine learning. Chemical Science, 9(2), 513-530. https://doi.org/10.1039/C7SC02664A",
    "Yue, X., Wang, Z., Huang, J., Parthasarathy, S., Moosavinasab, S., Huang, Y., Lin, S. M., Zhang, W., Zhang, P., & Sun, H. (2020). Graph embedding on biomedical networks: Methods, applications and evaluations. Bioinformatics, 36(4), 1241-1251. https://doi.org/10.1093/bioinformatics/btz718",
]


REVIEW = """# 2. REVISIÓN DE LITERATURA

La literatura reciente muestra una paradoja clara: existen muchos benchmarks y muchos modelos de grafos biomédicos, pero pocos instrumentos permiten saber si una mejora refleja aprendizaje biológico o una ventaja experimental inducida por el dataset. OGB formalizó una cultura de splits, métricas y loaders reproducibles para aprendizaje en grafos (Hu et al., 2020), y el benchmark de Dwivedi et al. reforzó que la comparación de GNNs exige presupuestos y protocolos comunes (Dwivedi et al., 2023). Sin embargo, ambos trabajos son transversales: su valor está en estandarizar graph ML, no en auditar de forma específica redes PPI humanas, solapamiento entre fuentes, negativos biomédicamente plausibles o calibración de probabilidades.

En biomedicina, TDC y MoleculeNet resolvieron una parte distinta del problema: convirtieron tareas terapéuticas y moleculares en datasets accesibles, comparables y reutilizables (Huang et al., 2022; Wu et al., 2018). Su debilidad para nuestro caso no es metodológica sino de alcance: privilegian moléculas, ADMET, farmacología o tareas de descubrimiento, mientras que BioGraphBench se concentra en redes proteína-proteína y función génica. OpenBioLink y Hetionet avanzan hacia grafos biomédicos heterogéneos y link prediction a gran escala (Breit et al., 2020; Himmelstein et al., 2017), pero su heterogeneidad también diluye una pregunta central: ¿qué tan difícil es predecir interacciones físicas PPI cuando se controlan splits, solapamientos y baselines estructurales simples?

OBNB es el antecedente más cercano para node classification biomédica, porque ofrece una caja de herramientas abierta para combinar redes biológicas y anotaciones de genes (Liu & Krishnan, 2024). Aun así, no está diseñado para responder simultáneamente link prediction PPI, clasificación funcional, ablaciones entre STRING/BioGRID y comparación calibrada de modelos no neuronales y GNNs. Esa brecha importa porque STRING y BioGRID son recursos abiertos y ampliamente usados, pero no son benchmarks por sí mismos: son bases de evidencia con sesgos de curación, cobertura desigual, redundancia y posible solapamiento (Oughtred et al., 2021; Szklarczyk et al., 2023). Usarlas sin una capa de auditoría puede convertir el benchmark en una medición de disponibilidad de datos, no de generalización.

Los trabajos de embeddings biomédicos y PPI confirman el riesgo. Yue et al. compararon múltiples métodos de embedding en redes biomédicas y mostraron que el desempeño depende fuertemente de tarea y fuente (Yue et al., 2020). Estudios específicos con GCN, GAT y GNNGL-PPI reportan mejoras prometedoras en PPI o variantes multicategoría de interacción proteica (Jha et al., 2022; Lv et al., 2024; Mohamed et al., 2022). No obstante, muchos de estos resultados se presentan como contribuciones de modelo: el protocolo de negativos, la separación estricta entre grafo de entrenamiento y grafo de prueba, la calibración y la comparación contra heurísticas fuertes no siempre quedan en el centro del argumento. Por eso una mejora de AUPRC puede ser difícil de interpretar si no se sabe cuánto explican grado, vecindad compartida o solapamiento entre recursos.

La literatura de evaluación refuerza esta preocupación. Shchur et al. mostraron que rankings de GNNs pueden cambiar por splits, hiperparámetros o prácticas de evaluación (Shchur et al., 2018), y node2vec recuerda que representaciones simples basadas en recorridos ya capturan regularidades topológicas fuertes (Grover & Leskovec, 2016). Las revisiones sobre GNNs en drug discovery y combinaciones sinérgicas muestran, además, un campo metodológicamente activo pero fragmentado: abundan arquitecturas, tareas y fuentes, mientras la discusión sobre comparabilidad, datos abiertos, calibración y baselines mínimos suele quedar dispersa. Esa dispersión produce un vacío práctico: hay suficiente evidencia para justificar GNNs biomédicas, pero no suficiente infraestructura para distinguir avance real de ventaja experimental.

En consecuencia, BioGraphBench entra como una propuesta de confianza antes que de complejidad: selecciona datasets abiertos, documenta usabilidad, define tareas PPI y funcionales, calcula features solo desde el grafo de entrenamiento, incluye ablaciones STRING/BioGRID, reporta calibración y establece baselines no neuronales exigentes. La contribución no es afirmar que una GNN gana, sino crear las condiciones para que una futura victoria sea científicamente interpretable."""


def write_csv() -> None:
    lines = ["id,authors,year,title,doi,url,open_preferred"]
    for row in REFERENCES:
        escaped = [str(x).replace('"', '""') for x in row]
        lines.append(",".join(f'"{x}"' for x in escaped))
    OUT_CSV.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)

    for line in REVIEW.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Referencias")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("2E74B5")

    for ref in APA_REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(ref)
        run.font.name = "Calibri"
        run.font.size = Pt(9)

    doc.save(OUT_DOCX)


def main() -> None:
    OUT_MD.write_text(REVIEW, encoding="utf-8")
    write_csv()
    build_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {OUT_CSV}")


if __name__ == "__main__":
    main()

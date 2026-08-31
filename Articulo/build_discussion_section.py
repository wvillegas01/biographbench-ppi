from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
OUT_MD = ARTICLE_DIR / "seccion_discusion.md"
OUT_DOCX = ARTICLE_DIR / "seccion_discusion_biographbench.docx"


DISCUSSION = """# 5. DISCUSIÓN

Los resultados de BioGraphBench muestran que el problema central no es únicamente entrenar mejores GNNs, sino construir condiciones bajo las cuales una mejora pueda interpretarse científicamente. En link prediction PPI, los modelos no neuronales basados en heurísticas estructurales y clasificadores supervisados alcanzaron desempeños muy altos, con HGB entre AUPRC=0,942 y 0,963. Esta observación cambia la lectura del benchmark: si una red PPI física puede predecirse con alta precisión usando grado, vecindad compartida y combinaciones supervisadas de esas señales, una arquitectura más compleja debe demostrar que aprende algo adicional. BioGraphBench no invalida las GNNs; eleva el estándar mínimo para defenderlas.

La comparación con trabajos no utilizados en la revisión de literatura refuerza esta interpretación. SkipGNN mostró que incorporar vecindarios de segundo orden puede mejorar la predicción de interacciones moleculares en redes incompletas (Huang et al., 2020), mientras BridgeDPI propuso nodos puente para explotar asociaciones proteína-proteína y droga-droga en predicción drug-protein (Wu et al., 2022). Estos resultados son coherentes con nuestro hallazgo de que la estructura relacional contiene señal predictiva fuerte; sin embargo, también sugieren una advertencia: si no se separa estrictamente el grafo de entrenamiento del grafo evaluado, las mejoras pueden confundirse con reutilización indirecta de conectividad. De forma complementaria, las revisiones de Sun et al. (2020), Feng et al. (2024) y Wan et al. (2024) describen un campo biomédico activo, con muchas arquitecturas y aplicaciones, pero menos consenso sobre calibración, comparabilidad y auditoría de datasets. BioGraphBench entra precisamente en ese punto: no como otro modelo, sino como una capa de evaluación que obliga a hacer explícita la procedencia, el split, los baselines y la confiabilidad probabilística.

Una contribución importante es que los resultados separan dos tipos de dificultad. Link prediction en PPI parece difícil desde el punto de vista biológico, pero empíricamente contiene una señal estructural muy explotable. En cambio, node classification sobre BioGRID+GOBP es más resistente: el GCN piloto apenas mejora Macro AUPRC frente al control constante y mantiene Micro-F1 bajo, con precisión cercana a 0,011. Esto sugiere que la propagación por la red no basta para recuperar funciones biológicas bajo desbalance extremo. BioGraphBench ofrece así un equilibrio útil: tareas donde los modelos deben superar baselines fuertes y tareas donde el desafío es extraer señal funcional débil sin inflar falsos positivos.

La calibración emerge como eje metodológico central. En link prediction, HGB no solo alcanza alto AUPRC, sino que mantiene ECE-10 muy bajo, entre 0,0037 y 0,0060. Esta diferencia importa porque muchos benchmarks se concentran en ranking, pero en biomedicina los scores suelen usarse para priorizar experimentos o hipótesis. Un modelo mal calibrado puede ordenar bien los pares y aun así inducir decisiones poco confiables. Por eso BioGraphBench debería conservar AUROC y AUPRC, pero tratarlas como métricas incompletas si no se acompañan de Brier score, ECE y estabilidad.

También hay una implicación editorial. Los resultados actuales no deben presentarse como una demostración de superioridad de BioGraphBench sobre benchmarks existentes ni como evidencia definitiva de que las GNNs fallan. La lectura más sólida es más precisa: el MVP demuestra que es posible construir un benchmark auditable con datasets abiertos, splits reproducibles, controles de leakage y baselines fuertes. Esa contribución es valiosa porque reduce el riesgo de sobreclaiming, una debilidad frecuente en estudios donde el énfasis recae en la arquitectura y no en el contrato experimental.

Finalmente, la principal limitación es que los resultados aún corresponden a una versión inicial con una semilla principal y un conjunto acotado de tareas. Para fortalecer el artículo, la siguiente versión debería incluir múltiples semillas, intervalos de confianza, pruebas pareadas y al menos una familia adicional de GNN bajo el mismo protocolo. Aun así, el potencial del trabajo es claro: BioGraphBench propone que la reproducibilidad no sea un apéndice del benchmark, sino su unidad básica de diseño."""


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)

    for line in DISCUSSION.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def main() -> None:
    OUT_MD.write_text(DISCUSSION, encoding="utf-8")
    build_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()

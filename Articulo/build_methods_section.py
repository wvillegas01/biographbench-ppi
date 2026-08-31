from __future__ import annotations

from pathlib import Path
import textwrap

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
FIG_DIR = ARTICLE_DIR / "figures"
OUT_MD = ARTICLE_DIR / "seccion_metodos.md"
OUT_DOCX = ARTICLE_DIR / "seccion_metodos_biographbench.docx"
METHOD_FIG = FIG_DIR / "figure_5_methodological_workflow.png"


def wrap_label(text: str, width: int = 25) -> str:
    return "\n".join(textwrap.wrap(text, width=width, break_long_words=False))


def build_method_figure() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.facecolor": "#ffffff",
            "figure.facecolor": "#ffffff",
            "font.size": 10,
        }
    )

    fig, ax = plt.subplots(figsize=(15.5, 7.6))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 7)
    ax.axis("off")

    boxes = [
        (0.45, 4.85, "Open datasets: STRING, BioGRID, OBNB", "#e8f1fb", "#2d6fa3"),
        (3.45, 4.85, "Audit layer: license, schema, leakage risk", "#edf7ed", "#3b7f43"),
        (6.45, 4.85, "Canonical graph: nodes, physical edges, labels", "#fff4dc", "#b97a1d"),
        (9.45, 4.85, "Task registry: LP and NC contracts", "#f4ecfb", "#7a4aa0"),
        (12.10, 4.85, "Reproducible outputs: splits, features, metrics", "#f1f3f5", "#5d6874"),
        (1.65, 2.00, "Leakage controls: resource-overlap ablations", "#fcebea", "#b94a48"),
        (4.95, 2.00, "Train-graph features: CN, Jaccard, AA, PA, degree", "#e8f1fb", "#2d6fa3"),
        (8.25, 2.00, "Baselines and GNN pilots: heuristics, HGB, GCN", "#edf7ed", "#3b7f43"),
        (11.55, 2.00, "Evaluation report: AUPRC, AUROC, ECE, F1", "#fff4dc", "#b97a1d"),
    ]

    for x, y, label, face, edge in boxes:
        patch = FancyBboxPatch(
            (x, y),
            2.35,
            1.02,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            linewidth=1.4,
            edgecolor=edge,
            facecolor=face,
        )
        ax.add_patch(patch)
        ax.text(
            x + 1.175,
            y + 0.51,
            wrap_label(label, 24),
            ha="center",
            va="center",
            color="#1f2933",
            fontsize=8.8,
            linespacing=1.15,
        )

    arrows = [
        ((2.80, 5.36), (3.45, 5.36)),
        ((5.80, 5.36), (6.45, 5.36)),
        ((8.80, 5.36), (9.45, 5.36)),
        ((11.80, 5.36), (12.10, 5.36)),
        ((7.10, 4.85), (2.82, 3.02)),
        ((4.00, 2.51), (4.95, 2.51)),
        ((7.30, 2.51), (8.25, 2.51)),
        ((10.60, 2.51), (11.55, 2.51)),
        ((12.72, 3.02), (13.10, 4.85)),
    ]
    for start, end in arrows:
        ax.add_patch(
            FancyArrowPatch(
                start,
                end,
                arrowstyle="-|>",
                mutation_scale=14,
                linewidth=1.2,
                color="#6b7280",
                connectionstyle="arc3,rad=0.0",
            )
        )

    ax.text(
        7.5,
        6.55,
        "BioGraphBench methodological workflow",
        ha="center",
        va="center",
        fontsize=15,
        weight="bold",
        color="#111827",
    )
    ax.text(
        7.5,
        0.72,
        "Every reported score must be traceable to an audited source, a canonical graph, a fixed split, train-only features, and a versioned metric report.",
        ha="center",
        va="center",
        fontsize=10,
        color="#374151",
    )
    fig.savefig(METHOD_FIG, dpi=300, bbox_inches="tight")
    plt.close(fig)


def build_markdown() -> str:
    md = r"""# 3. METODOLOGÍA

## 3.1. Diseño general del benchmark y unidad experimental

BioGraphBench se diseñó como un benchmark reproducible para evaluar aprendizaje sobre grafos biomédicos bajo una restricción central: ningún resultado debe depender de decisiones implícitas de descarga, filtrado, particionado o construcción de features. Por esta razón, el método no inicia con arquitecturas de modelos, sino con una representación auditable de cada tarea. Cada dataset se transforma en un objeto experimental compuesto por nodos, aristas, atributos, etiquetas, metadatos de procedencia y particiones versionadas:

$$
\mathcal{B}_k=\left(\mathcal{G}_k,\mathcal{T}_k,\mathcal{S}_k,\Phi_k,\mathcal{M}_k\right), \qquad
\mathcal{G}_k=(V_k,E_k,X_k,Y_k).
\tag{1}
$$

En la Ecuación (1), $\mathcal{B}_k$ representa la instancia benchmark para el dataset $k$; $\mathcal{G}_k$ es el grafo canónico; $\mathcal{T}_k$ define la tarea; $\mathcal{S}_k$ contiene los splits; $\Phi_k$ define las funciones de características, y $\mathcal{M}_k$ define las métricas. Esta formulación obliga a que cualquier comparación futura, incluyendo GNNs, declare exactamente qué grafo, qué tarea, qué split y qué métricas está usando.

La unidad mínima de evaluación se define de forma distinta según la familia de tarea. Para predicción de enlaces, la unidad es un par no dirigido de proteínas $(u,v)$; para clasificación de nodos, la unidad es una proteína $v$ con un vector multilabel de funciones biológicas. Formalmente:

$$
z_i =
\begin{cases}
(u_i,v_i,y_i), & \text{si } \mathcal{T}_k \text{ es predicción de enlaces},\\
(v_i,\mathbf{y}_i), & \text{si } \mathcal{T}_k \text{ es clasificación multilabel de nodos}.
\end{cases}
\tag{2}
$$

Esta definición separa claramente el objeto biológico observado de la tarea estadística que se construye sobre él. El objetivo no es maximizar una métrica aislada, sino asegurar que cada $z_i$ sea reconstruible desde los artefactos públicos del benchmark.

![Figure 5](figures/figure_5_methodological_workflow.png)

**Figure 5. Methodological workflow of BioGraphBench.** The pipeline starts from open biological resources, applies audit and canonicalization rules, defines task contracts and leakage controls, computes train-graph features, evaluates baselines and produces versioned reports.

La Figura 5 resume el flujo metodológico. El elemento crítico es que los modelos aparecen tarde en el pipeline: antes se auditan fuentes, se normaliza el grafo, se define la tarea, se controla leakage y se fijan features calculables solo desde el grafo de entrenamiento. Esta decisión reduce el riesgo de que el benchmark mida disponibilidad accidental de datos o solapamiento entre recursos en lugar de aprendizaje estructural.

## 3.2. Selección, auditoría y canonicalización de datasets

Los datasets candidatos se evaluaron mediante una matriz de auditabilidad que considera apertura, trazabilidad, descarga completa, licencia, esquema, identificadores y posibilidad de reconstrucción. Un recurso se acepta únicamente si satisface simultáneamente criterios mínimos de disponibilidad, claridad semántica y utilidad experimental:

$$
A(d)=\mathbb{1}\left[
L_d=1 \land R_d=1 \land C_d=1 \land U_d=1
\right],
\tag{3}
$$

donde $L_d$ indica licencia o términos de uso compatibles, $R_d$ reproducibilidad de descarga, $C_d$ claridad del esquema y $U_d$ usabilidad para una tarea de grafo. Esta regla evita incluir datasets que parezcan atractivos pero no puedan auditarse o descargarse completamente.

Después de seleccionar un recurso, las aristas se convierten a una forma canónica no dirigida. Para toda interacción física entre proteínas $u$ y $v$, BioGraphBench registra una sola arista ordenada lexicográficamente:

$$
\operatorname{canon}(u,v)=\left(\min(u,v),\max(u,v)\right), \qquad
E=\{\operatorname{canon}(u,v):(u,v)\in E_{\mathrm{raw}},u\neq v\}.
\tag{4}
$$

La Ecuación (4) elimina duplicados inducidos por orden, autointeracciones y registros redundantes. En STRING y BioGRID se retienen interacciones físicas humanas; en OBNB se conserva la relación entre proteínas y anotaciones GOBP cuando estas pueden asociarse a nodos presentes en la red PPI. Esta decisión hace que las tareas de link prediction y node classification compartan un espacio de nodos comparable.

## 3.3. Definición de tareas, particionado y control de leakage

Para predicción de enlaces, las aristas positivas se dividen en entrenamiento, validación y prueba mediante una partición disjunta y reproducible:

$$
E^+=E^+_{\mathrm{train}}\dot{\cup}E^+_{\mathrm{val}}\dot{\cup}E^+_{\mathrm{test}}, \qquad
E^+_a\cap E^+_b=\varnothing \; \forall a\neq b.
\tag{5}
$$

El grafo usado para calcular features y entrenar modelos contiene únicamente $E^+_{\mathrm{train}}$. Esto impide que una feature estructural use vecindarios formados por aristas de validación o prueba. Los negativos se muestrean desde el complemento del grafo observado, excluyendo pares positivos y preservando una proporción controlada:

$$
E^-_s \subseteq \binom{V}{2}\setminus E^+, \qquad
\rho_s=\frac{|E^-_s|}{|E^+_s|}, \quad s\in\{\mathrm{train},\mathrm{val},\mathrm{test}\}.
\tag{6}
$$

Las tareas de ablación STRING/BioGRID agregan un segundo control: remover del conjunto evaluado aquellos pares que aparecen como positivos en el otro recurso. Si $E^{(a)}$ y $E^{(b)}$ son dos fuentes PPI, la variante sin solapamiento se define como:

$$
E^{(a\setminus b)}_{\mathrm{eval}}=\{e\in E^{(a)}_{\mathrm{eval}}:e\notin E^{(b)}\}.
\tag{7}
$$

La Ecuación (7) no afirma independencia biológica total entre recursos, pero sí reduce una forma concreta de leakage por reutilización de interacciones ya registradas en otra base. En node classification, el split se realiza sobre nodos y las etiquetas GOBP se mantienen fuera de las features estructurales. Así, la evaluación pregunta si la estructura PPI permite recuperar funciones biológicas sin usar la etiqueta como atributo.

## 3.4. Extracción de features, baselines y modelos

Las features de predicción de enlaces se calculan exclusivamente sobre el grafo de entrenamiento $G_{\mathrm{train}}=(V,E^+_{\mathrm{train}})$. Para un par $(u,v)$ se define el vector estructural:

$$
\phi(u,v)=\left[
\mathrm{CN}(u,v),\mathrm{Jaccard}(u,v),\mathrm{AA}(u,v),
\mathrm{PA}(u,v),\deg(u),\deg(v)
\right],
\quad
\mathrm{CN}=|\Gamma(u)\cap\Gamma(v)|,\quad
\mathrm{Jaccard}=\frac{|\Gamma(u)\cap\Gamma(v)|}{|\Gamma(u)\cup\Gamma(v)|},
\quad
\mathrm{AA}=\sum_{w\in\Gamma(u)\cap\Gamma(v)}\frac{1}{\log(\deg(w)+\epsilon)},\quad
\mathrm{PA}=\deg(u)\deg(v).
\tag{8}
$$

Estas features sirven para dos tipos de modelos: heurísticas puras, donde el score es directamente una función estructural, y modelos supervisados, donde $\phi(u,v)$ alimenta clasificadores como Logistic Regression, Random Forest e HistGradientBoosting. Para link prediction y node classification multilabel, las pérdidas se definen como:

$$
\mathcal{L}_{\mathrm{LP}}(\theta)=
-\frac{1}{n}\sum_{i=1}^{n}
\left[y_i\log \hat{p}_\theta(u_i,v_i)+(1-y_i)\log(1-\hat{p}_\theta(u_i,v_i))\right],
\quad
\mathcal{L}_{\mathrm{NC}}(\theta)=
-\frac{1}{|V_{\mathrm{train}}|C}
\sum_{v\in V_{\mathrm{train}}}\sum_{c=1}^{C}
\left[y_{vc}\log \hat{p}_{vc}+(1-y_{vc})\log(1-\hat{p}_{vc})\right].
\tag{9}
$$

**Algoritmo 1. Protocolo reproducible de BioGraphBench.**

```text
Input:
  Datasets candidatos D, semilla r, razones de split, métricas M
Output:
  Reportes reproducibles con tareas, splits, features, modelos y métricas

1. Para cada dataset d en D:
2.   Auditar licencia, descarga, esquema, identificadores y cobertura.
3.   Si A(d)=0, excluir d y registrar la causa.
4.   Canonicalizar nodos y aristas; remover duplicados y autointeracciones.
5.   Definir contratos de tarea para link prediction o node classification.
6.   Construir splits reproducibles usando la semilla r.
7.   Para link prediction, muestrear negativos desde el complemento del grafo.
8.   Aplicar controles de leakage y ablaciones de solapamiento cuando corresponda.
9.   Calcular features únicamente desde el grafo de entrenamiento.
10.  Entrenar heurísticas, baselines supervisados y modelos piloto.
11.  Evaluar AUROC, AUPRC, calibración, F1, precision y recall según la tarea.
12.  Exportar matrices, artefactos intermedios, configuración y reportes.
```

## 3.5. Métricas, calibración y garantía de reproducibilidad

La evaluación combina métricas de ranking, métricas de decisión, calibración y trazabilidad. En link prediction se priorizan AUROC y AUPRC, dado que AUPRC es más informativa bajo desbalance entre pares observados y no observados. En node classification se reportan Macro AUROC, Macro AUPRC y métricas microagregadas después de ajustar umbrales en validación. El bloque final de evaluación se define como:

$$
\mathcal{M}=
\{\mathrm{AUROC},\mathrm{AUPRC},\mathrm{ECE},\mathrm{Brier},\mathrm{MicroF1}\},
\quad
\mathrm{ECE}=\sum_{b=1}^{B}\frac{|I_b|}{n}
\left|\operatorname{acc}(I_b)-\operatorname{conf}(I_b)\right|,
\quad
R=f(D_{\mathrm{raw}},C_{\mathrm{audit}},S_r,\Phi,\Theta,\mathcal{M}),
\tag{10}
$$

donde $D_{\mathrm{raw}}$ son los datos originales, $C_{\mathrm{audit}}$ las reglas de auditoría, $S_r$ los splits generados con semilla $r$, $\Phi$ las features, $\Theta$ la configuración de modelos y $\mathcal{M}$ las métricas. Esta definición permite auditar no solo el valor final de una métrica, sino también el camino que llevó a ella.
"""

    OUT_MD.write_text(md, encoding="utf-8")
    return md


def add_code_block(doc: Document, code: str) -> None:
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_equation(doc: Document, equation_lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(" ".join(line.strip() for line in equation_lines))
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(17, 24, 39)


def build_docx(md: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)

    in_code = False
    code_lines: list[str] = []
    equation_lines: list[str] = []
    in_equation = False

    for raw in md.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code = False
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line == "$$":
            if not in_equation:
                in_equation = True
                equation_lines = []
            else:
                add_equation(doc, equation_lines)
                in_equation = False
            continue

        if in_equation:
            equation_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("![Figure 5]"):
            doc.add_picture(str(METHOD_FIG), width=Inches(6.9))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line.strip("*"))
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 77, 120)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(line.replace("`", ""))
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def main() -> None:
    build_method_figure()
    md = build_markdown()
    build_docx(md)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {METHOD_FIG}")


if __name__ == "__main__":
    main()

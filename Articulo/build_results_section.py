from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
REPORTS_DIR = ROOT / "reports"
FIG_DIR = ARTICLE_DIR / "figures"
OUT_MD = ARTICLE_DIR / "seccion_resultados.md"
OUT_DOCX = ARTICLE_DIR / "seccion_resultados_biographbench.docx"


def read_csv_rows(path: Path):
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value, digits=3):
    if value == "" or value is None:
        return ""
    return f"{float(value):.{digits}f}"


def load_tables():
    task_rows = read_csv_rows(REPORTS_DIR / "task_definition_matrix.csv")
    baseline_rows = read_csv_rows(REPORTS_DIR / "model_baseline_matrix.csv")
    supervised = [r for r in read_json(REPORTS_DIR / "link_prediction_supervised.json") if r["split"] == "test"]
    node_tuned = [r for r in read_json(REPORTS_DIR / "node_classification_threshold_tuning.json") if r["split"] == "test"]
    node_logreg = [r for r in read_json(REPORTS_DIR / "node_classification_logreg.json") if r["split"] == "test"]
    return task_rows, baseline_rows, supervised, node_tuned, node_logreg


def markdown_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(x) for x in row) + " |")
    return "\n".join(lines)


def build_markdown():
    task_rows, baseline_rows, supervised, node_tuned, node_logreg = load_tables()

    task_table_rows = []
    for r in task_rows:
        task_table_rows.append(
            [
                f"`{r['task_id']}`",
                f"`{r['dataset_id']}`",
                r["task_family"],
                r["task_type"],
                r["primary_metrics"].replace(";", ", "),
                r["paper_role"],
            ]
        )

    lp_datasets = [
        "string_human_physical_v12",
        "biogrid_human_physical",
        "biogrid_human_physical_no_string_overlap",
        "string_human_physical_no_biogrid_overlap",
    ]
    labels = {
        "string_human_physical_v12": "STRING",
        "biogrid_human_physical": "BioGRID",
        "biogrid_human_physical_no_string_overlap": "BioGRID no STRING",
        "string_human_physical_no_biogrid_overlap": "STRING no BioGRID",
    }
    lp_rows = []
    for dataset in lp_datasets:
        best_h = max(
            [
                r
                for r in baseline_rows
                if r["task"] == "link_prediction"
                and r["dataset"] == dataset
                and r["model"] in {"common_neighbors", "jaccard", "adamic_adar", "preferential_attachment"}
            ],
            key=lambda r: float(r["auprc"]),
        )
        hgb = next(
            r
            for r in baseline_rows
            if r["task"] == "link_prediction" and r["dataset"] == dataset and r["model"] == "hist_gradient_boosting"
        )
        lp_rows.append(
            [
                labels[dataset],
                best_h["model"],
                fmt(best_h["auprc"]),
                fmt(hgb["auroc"]),
                fmt(hgb["auprc"]),
                f"+{float(hgb['auprc']) - float(best_h['auprc']):.3f}",
            ]
        )

    cal_rows = []
    for dataset in lp_datasets:
        for model in ["logistic_regression", "random_forest", "hist_gradient_boosting"]:
            row = next(r for r in supervised if r["dataset"] == dataset and r["model"] == model)
            cal_rows.append(
                [
                    labels[dataset],
                    model,
                    fmt(row["auprc"]),
                    fmt(row["brier"], 4),
                    fmt(row["ece_10"], 4),
                    fmt(row["train_seconds"], 2),
                ]
            )

    node_rows = []
    constant = next(r for r in node_logreg if r["feature"] == "constant")
    node_rows.append(["constant_logreg", "constant", fmt(constant["macro_auroc"]), fmt(constant["macro_auprc"]), fmt(constant["micro_f1"]), "", ""])
    for model in ["logistic_regression", "mlp", "gcn"]:
        row = next(r for r in node_tuned if r["model"] == model)
        node_rows.append(
            [
                model,
                row["feature"],
                fmt(row["macro_auroc"]),
                fmt(row["macro_auprc"]),
                fmt(row["micro_f1"]),
                fmt(row["micro_precision"]),
                fmt(row["micro_recall"]),
            ]
        )

    criteria_rows = [
        [
            "Link prediction PPI",
            "Superar HGB, no solo heurísticas",
            "HGB alcanza AUPRC 0,942-0,963; la mejor heurística llega hasta 0,934",
            "Una GNN debe mejorar AUPRC y calibración, no solo AUROC",
        ],
        [
            "Ablaciones STRING/BioGRID",
            "Mantener ganancia tras remover solapamiento",
            "BioGRID no STRING conserva AUPRC HGB=0,942; STRING no BioGRID llega a 0,963",
            "La mejora debe persistir sin depender de pares compartidos entre recursos",
        ],
        [
            "Calibración",
            "Reportar probabilidades confiables",
            "HGB mantiene ECE-10 entre 0,0037 y 0,0060; LogReg llega hasta 0,0601",
            "Modelos futuros deben reportar ECE/Brier junto a AUPRC",
        ],
        [
            "Node classification",
            "Resolver desbalance multilabel",
            "GCN mejora Macro AUPRC a 0,016 pero Micro-F1 queda en 0,021",
            "La contribución real requiere mejorar AUPRC sin colapsar precision",
        ],
        [
            "Reproducibilidad",
            "Estabilidad con múltiples semillas",
            "MVP usa semilla 42; splits y features ya son auditables",
            "La versión final debe reportar media, desviación y pruebas pareadas",
        ],
    ]

    figure_paths = {
        "fig1": "figures/figure_1_supervised_metric_profiles.png",
        "fig2": "figures/figure_2_baseline_distribution_violin_box.png",
        "fig3": "figures/figure_3_gain_and_calibration_panels.png",
        "fig4": "figures/figure_4_node_classification_dual_panels.png",
    }

    md = f"""# 4. RESULTADOS

## 4.1. Tareas aceptadas y alcance empírico del benchmark

BioGraphBench fue diseñado bajo una premisa distinta a la de un benchmark centrado exclusivamente en maximizar desempeño de modelos: antes de comparar arquitecturas, el benchmark debe demostrar que sus tareas son auditables, reproducibles y resistentes a leakage. Bajo este criterio, el MVP retiene cinco tareas: cuatro de predicción de enlaces en redes PPI físicas y una de clasificación multilabel de nodos sobre funciones biológicas GOBP. La Tabla 1 resume el contrato experimental de estas tareas y separa explícitamente las tareas principales de las tareas de ablación.

**Tabla 1. Tareas aceptadas en el MVP de BioGraphBench.**

{markdown_table(["Task ID", "Dataset", "Familia", "Tipo", "Métricas primarias", "Rol"], task_table_rows)}

El resultado metodológico más importante de la Tabla 1 no es el número de tareas, sino su heterogeneidad controlada. Las dos tareas principales de link prediction (`STRING` y `BioGRID`) permiten evaluar aprendizaje estructural sobre redes PPI físicas; las dos variantes sin solapamiento permiten probar si los resultados sobreviven cuando se remueven pares compartidos entre recursos; y `obnb_biogrid_gobp` introduce una tarea de node classification mucho más desbalanceada y menos trivial. Así, el benchmark no depende de una única lectura de desempeño: combina señal topológica fuerte, controles de independencia parcial y una tarea funcional difícil.

## 4.2. Predicción de enlaces en redes PPI: heurísticas fuertes y modelos supervisados

La primera pregunta empírica es si las tareas PPI son suficientemente difíciles como para justificar modelos más complejos. Los resultados muestran una respuesta matizada: las heurísticas clásicas son sorprendentemente fuertes, pero los modelos supervisados sobre heurísticas de pares todavía añaden ganancia consistente. Esto es crucial para BioGraphBench porque establece una vara mínima: cualquier GNN futura debe superar no solo baselines triviales, sino también heurísticas estructurales bien calibradas.

**Tabla 2. Mejor heurística clásica versus mejor baseline supervisado en link prediction.**

{markdown_table(["Dataset", "Mejor heurística", "AUPRC heurística", "AUROC HGB", "AUPRC HGB", "Ganancia AUPRC"], lp_rows)}

En la Tabla 2, HistGradientBoosting (HGB) alcanza AUPRC entre 0,942 y 0,963 en las cuatro tareas PPI. La magnitud de estos valores no debe interpretarse como evidencia de que la tarea esté resuelta biológicamente. Más bien, indica que las redes PPI físicas poseen regularidades topológicas muy fuertes: nodos con vecindarios compartidos, conectividad preferencial y concentración de grado ya explican gran parte de la separabilidad entre pares positivos y negativos. El caso más ilustrativo es `biogrid_human_physical_no_string_overlap`, donde Preferential Attachment alcanza AUPRC=0,934 y HGB solo mejora a 0,942. En esta variante, remover solapamiento directo con STRING no destruye la señal topológica; al contrario, revela que parte sustancial de la predictibilidad proviene de la organización global de la red.

![Figure 1](figures/figure_1_supervised_metric_profiles.png)

**Figure 1. Supervised link-prediction metric profiles across accepted PPI tasks.** Panel A reports test AUROC and Panel B reports test AUPRC for logistic regression, random forest and HistGradientBoosting trained on train-graph pair heuristics.

La Figura 1 muestra que los perfiles de AUROC y AUPRC son paralelos para Random Forest y HGB, con diferencias pequeñas pero sistemáticas a favor de HGB. En AUROC, HGB se mueve en un rango estrecho de 0,941 a 0,955, mientras que Random Forest queda prácticamente superpuesto entre 0,941 y 0,955. En AUPRC ocurre algo similar: HGB oscila entre 0,942 y 0,963, y Random Forest entre 0,941 y 0,963. La lectura importante no es que HGB gane por amplio margen, sino que el techo de desempeño no neuronal ya es alto y estable. Logistic regression queda por debajo en todas las tareas, con caídas más visibles en BioGRID y BioGRID sin STRING, donde su AUPRC baja a 0,935 y 0,933, respectivamente. Esto sugiere que una combinación lineal de heurísticas no captura completamente interacciones no lineales entre grado, vecindad compartida y conectividad local. Para BioGraphBench, esta figura define el umbral empírico mínimo: una GNN que no supere aproximadamente AUPRC=0,94-0,96 en PPI no estaría aportando evidencia fuerte de aprendizaje adicional.

![Figure 2](figures/figure_2_baseline_distribution_violin_box.png)

**Figure 2. Baseline score distributions across the accepted PPI tasks.** Panel A summarizes test AUROC and Panel B summarizes test AUPRC using violin distributions, box summaries and task-level points for classical heuristics and supervised pair-heuristic models.

La Figura 2 permite observar algo que la tabla oculta: no todas las heurísticas fallan del mismo modo ni tienen la misma estabilidad entre datasets. Jaccard es el caso más débil y variable: su AUPRC cae hasta 0,726 en BioGRID sin STRING, a pesar de mantenerse por encima de 0,90 en STRING y STRING sin BioGRID. Esta sensibilidad indica que normalizar por la unión de vecindarios penaliza con fuerza redes donde el patrón dominante no es vecindad local compartida sino conectividad por grado. Adamic-Adar, por el contrario, se mantiene en una banda más estable, con AUPRC entre 0,878 y 0,927; Preferential Attachment alcanza 0,934 en BioGRID sin STRING y supera a Adamic-Adar en las tareas BioGRID, lo que revela una señal de grado particularmente fuerte en esa red. Los modelos supervisados forman una distribución compacta en la región superior: Logistic Regression, Random Forest y HGB concentran sus valores por encima de 0,932. Esta compactación es científicamente relevante porque muestra que la supervisión no solo eleva el desempeño, sino que reduce la dependencia del baseline respecto al sesgo estructural específico de cada red.

## 4.3. Ganancia supervisada, calibración y confiabilidad de los baselines

La segunda pregunta empírica es si la ganancia de los modelos supervisados representa una mejora sustantiva o simplemente una redistribución marginal de scores. Para responderlo, comparamos el mejor baseline clásico de cada tarea contra HGB y analizamos simultáneamente calibración mediante ECE-10, Brier score y tiempo de entrenamiento.

**Tabla 3. Calibración y costo de modelos supervisados en link prediction.**

{markdown_table(["Dataset", "Modelo", "AUPRC", "Brier", "ECE-10", "Train s"], cal_rows)}

La Tabla 3 muestra un patrón consistente: HGB y Random Forest obtienen AUPRC similares, pero HGB tiende a entrenar mucho más rápido y a mantener ECE-10 bajo. En `STRING`, por ejemplo, Random Forest alcanza AUPRC=0,950 con 35,59 s de entrenamiento, mientras que HGB alcanza AUPRC=0,951 con 5,32 s. En `BioGRID`, la diferencia de tiempo es aún mayor: 53,20 s frente a 7,41 s. Esto posiciona a HGB como baseline supervisado preferente para el MVP: no necesariamente porque domine todas las métricas por amplio margen, sino porque ofrece buen desempeño, calibración razonable y costo computacional menor.

![Figure 3](figures/figure_3_gain_and_calibration_panels.png)

**Figure 3. Added value and reliability of supervised pair-heuristic baselines.** Panel A shows the paired AUPRC gain from the best classical heuristic to HGB for each PPI task. Panel B shows test AUPRC versus ECE-10, with marker size proportional to training time.

La Figura 3(a) cuantifica la ganancia de pasar de la mejor heurística clásica a HGB. La ganancia es mayor en `STRING no BioGRID` (+0,042), seguida de BioGRID (+0,025) y STRING (+0,024), mientras que cae a +0,008 en `BioGRID no STRING`. Esta asimetría es informativa: cuando una heurística como Preferential Attachment ya captura la estructura dominante, el margen para supervisión adicional disminuye; cuando la estructura local no basta, el modelo supervisado combina señales de forma más efectiva. En otras palabras, la ganancia supervisada no es uniforme; depende de qué componente topológico domina la red después del filtrado. La Figura 3(b) añade una dimensión de confiabilidad. Logistic Regression puede obtener AUPRC competitivo en `STRING no BioGRID` (0,958), pero su ECE-10 es 0,0286, casi cinco veces mayor que HGB (0,0060). En BioGRID sin STRING, el contraste es aún más fuerte: Logistic Regression alcanza ECE-10=0,0601, mientras HGB queda en 0,0039. Esto significa que un modelo puede ordenar pares relativamente bien y aun así producir probabilidades poco confiables. Por esta razón, BioGraphBench debe reportar calibración junto con ranking: en aplicaciones biomédicas downstream, un score mal calibrado puede inducir priorizaciones experimentales engañosas.

## 4.4. Clasificación multilabel de nodos: una tarea difícil y desbalanceada

La tarea `obnb_biogrid_gobp` cuenta una historia diferente a link prediction. Aquí, features estructurales simples no bastan para producir desempeño fuerte, y el GCN piloto tampoco resuelve el problema. La Tabla 4 resume los resultados de clasificación multilabel usando features constantes, degree-bin logistic regression, MLP y GCN con threshold tuning por tarea seleccionado sobre validación.

**Tabla 4. Resultados de node classification en OBNB BioGRID+GOBP.**

{markdown_table(["Modelo", "Feature", "Macro AUROC", "Macro AUPRC", "Micro-F1", "Precision", "Recall"], node_rows)}

La Tabla 4 muestra que `one_hot_log_degree` mejora el control constante en Macro AUROC, pero la mejora no transforma la tarea en un problema fácil. Logistic regression con bins de grado alcanza Macro AUROC=0,531 y Macro AUPRC=0,014, apenas por encima del control constante en AUPRC. El GCN alcanza la mejor Macro AUPRC del conjunto (0,016), pero su Micro-F1 threshold-tuned queda en 0,021. Esta divergencia entre ranking y decisión revela un problema central: el modelo puede ordenar algunos labels raros ligeramente mejor, pero convertir scores en decisiones multilabel confiables sigue siendo difícil.

![Figure 4](figures/figure_4_node_classification_dual_panels.png)

**Figure 4. OBNB BioGRID+GOBP node-classification trade-offs.** Panel A compares discrimination metrics, Macro AUROC and Macro AUPRC. Panel B compares threshold-tuned Micro-F1, precision and recall.

La Figura 4(a) muestra que Macro AUROC y Macro AUPRC no seleccionan necesariamente el mismo modelo: logistic regression con degree bins maximiza AUROC (0,531), mientras que GCN maximiza AUPRC (0,016). La diferencia parece pequeña en escala absoluta, pero es relevante porque el control constante ya tiene Macro AUPRC=0,011; por tanto, el GCN mejora aproximadamente 49% sobre el control en AUPRC relativa, aunque sigue lejos de una tarea resuelta. El MLP, pese a ser más flexible que logistic regression, cae a Macro AUROC=0,488 y Macro AUPRC=0,013, lo que sugiere que mayor capacidad no garantiza mejor señal bajo features estructurales tan simples. La Figura 4(b) revela el costo del threshold tuning: MLP y GCN alcanzan recall alto (0,688 y 0,605), pero con precision extremadamente baja (0,011 en ambos casos), lo que mantiene Micro-F1 en 0,022 y 0,021. Logistic Regression presenta un equilibrio menos agresivo: recall=0,416, precision=0,013 y Micro-F1=0,025, el mejor F1 de la tabla. En otras palabras, el problema no es solo aprender embeddings o propagar información por la red; también es calibrar decisiones multilabel bajo prevalencias muy bajas sin convertir el modelo en un predictor excesivamente positivo.

## 4.5. Implicaciones empíricas para futuras GNNs en BioGraphBench

Los resultados anteriores permiten transformar BioGraphBench de una colección de datasets en un instrumento de evaluación con criterios empíricos explícitos. La implicación principal es que el benchmark no debe premiar complejidad arquitectónica por sí misma. En link prediction, la topología de PPI es tan informativa que modelos no neuronales ya alcanzan AUPRC superior a 0,94 en todas las tareas aceptadas. Por tanto, una GNN futura tendría que demostrar una mejora sobre HGB y Random Forest, no simplemente superar heurísticas débiles. Además, tendría que hacerlo conservando calibración comparable: HGB mantiene ECE-10 entre 0,0037 y 0,0060, mientras Logistic Regression llega hasta 0,0601. Esto define un criterio doble: desempeño de ranking y confiabilidad probabilística.

En node classification, el criterio es distinto. La tarea `obnb_biogrid_gobp` no exige superar un baseline fuerte de AUPRC alto, sino resolver un problema de desbalance donde los modelos actuales apenas logran Macro AUPRC entre 0,013 y 0,016 con features estructurales. Aquí, una GNN útil no debería evaluarse solo por AUROC; debería mejorar Macro AUPRC, Micro-F1 y precision simultáneamente. Un modelo que aumente recall pero mantenga precision alrededor de 0,011 estaría recuperando muchas etiquetas positivas a costa de demasiados falsos positivos, lo que en un contexto de anotación funcional limitaría su utilidad práctica.

**Tabla 5. Criterios empíricos derivados para modelos futuros en BioGraphBench.**

{markdown_table(["Familia", "Criterio", "Evidencia actual", "Implicación"], criteria_rows)}

La Tabla 5 resume el valor práctico de esta sección de resultados. BioGraphBench no solo reporta números: convierte esos números en condiciones de evaluación. Para link prediction, el umbral mínimo no es "mejor que azar", sino mejor que HGB sobre heurísticas estructurales de entrenamiento. Para las ablaciones, el requisito es mostrar que la mejora no depende del solapamiento STRING-BioGRID. Para node classification, el desafío es construir modelos que mejoren ranking y decisión multilabel sin explotar únicamente el grado. Finalmente, para una versión publicable más fuerte, estos criterios deben repetirse con múltiples semillas, intervalos de confianza y pruebas pareadas. Esta es la diferencia entre presentar resultados exploratorios y proponer un benchmark reproducible capaz de sostener comparaciones científicas.
"""
    OUT_MD.write_text(md, encoding="utf-8")
    return md


def set_run_style(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph_from_markdown(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_style(run, 11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_style(run, 9, italic=True, color="555555")
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_style(run, 8.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value).replace("`", "")
            for p in cells[i].paragraphs:
                for run in p.runs:
                    set_run_style(run, 8)
    doc.add_paragraph()
    return table


def build_docx(md: str):
    task_rows, baseline_rows, supervised, node_tuned, node_logreg = load_tables()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.color.rgb = RGBColor.from_string(color)

    lines = md.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line:
            idx += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("**Tabla "):
            add_paragraph_from_markdown(doc, line.replace("**", ""))
            idx += 2
            headers = [h.strip() for h in lines[idx].strip("|").split("|")]
            idx += 2
            rows = []
            while idx < len(lines) and lines[idx].startswith("|"):
                rows.append([c.strip() for c in lines[idx].strip("|").split("|")])
                idx += 1
            add_table(doc, headers, rows)
            continue
        elif line.startswith("![Figure"):
            image_rel = line.split("(", 1)[1].split(")", 1)[0]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(ARTICLE_DIR / image_rel), width=Inches(6.5))
        elif line.startswith("**Figure "):
            add_caption(doc, line.replace("**", ""))
        else:
            add_paragraph_from_markdown(doc, line.replace("`", ""))
        idx += 1

    doc.save(OUT_DOCX)


def main() -> int:
    md = build_markdown()
    build_docx(md)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

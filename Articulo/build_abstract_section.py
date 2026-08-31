from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
OUT_MD = ARTICLE_DIR / "titulo_resumen_keywords.md"
OUT_DOCX = ARTICLE_DIR / "titulo_resumen_keywords_biographbench.docx"


TITLE = "BioGraphBench: un benchmark reproducible y auditable para aprendizaje sobre grafos biomédicos"

ABSTRACT = """Introducción: El aprendizaje sobre grafos biomédicos ha impulsado el uso de GNNs en predicción de interacciones proteína-proteína y anotación funcional; sin embargo, muchos resultados dependen de datasets, splits y baselines poco auditables. Objetivo: Este trabajo presenta BioGraphBench, un benchmark diseñado para priorizar confianza experimental antes que complejidad arquitectónica y facilitar comparaciones transparentes. Metodología: Se construyó un pipeline reproducible que audita recursos abiertos, canonicaliza redes STRING y BioGRID, define tareas de link prediction y node classification, genera splits controlados, calcula features solo desde el grafo de entrenamiento y evalúa heurísticas, modelos supervisados y pilotos neuronales mediante AUROC, AUPRC, calibración y métricas multilabel. Resultados: En link prediction PPI, HistGradientBoosting alcanzó AUPRC entre 0,942 y 0,963, con ECE-10 entre 0,0037 y 0,0060, estableciendo un baseline no neuronal exigente. En node classification BioGRID+GOBP, el GCN piloto obtuvo Macro AUPRC=0,016 y Micro-F1 cercano a 0,021, evidenciando una tarea funcional aún difícil bajo desbalance extremo. Conclusión: BioGraphBench muestra que un benchmark biomédico serio debe iniciar con trazabilidad, controles de leakage, baselines fuertes y métricas calibradas. Su contribución es una base reproducible para evaluar futuras GNNs bajo condiciones científicamente interpretables y reconstruibles desde los datos originales."""

KEYWORDS = [
    "BioGraphBench",
    "grafos biomédicos",
    "benchmark reproducible",
    "predicción de enlaces",
    "redes proteína-proteína",
]


def build_markdown() -> str:
    md = (
        f"# {TITLE}\n\n"
        f"## Resumen\n\n{ABSTRACT}\n\n"
        f"**Palabras clave:** {', '.join(KEYWORDS)}.\n"
    )
    OUT_MD.write_text(md, encoding="utf-8")
    return md


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("2E74B5")

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run("Resumen")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("2E74B5")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(ABSTRACT)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)

    kw = doc.add_paragraph()
    kw.paragraph_format.space_after = Pt(4)
    r = kw.add_run("Palabras clave: ")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r2 = kw.add_run(", ".join(KEYWORDS) + ".")
    r2.font.name = "Calibri"
    r2.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def main() -> None:
    build_markdown()
    build_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()

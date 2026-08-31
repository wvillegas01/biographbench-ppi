from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTICLE_DIR = ROOT / "Articulo"
OUT_MD = ARTICLE_DIR / "seccion_introduccion.md"
OUT_DOCX = ARTICLE_DIR / "seccion_introduccion_biographbench.docx"


INTRODUCTION = """# 1. INTRODUCCIÓN

El aprendizaje sobre grafos biomédicos se ha convertido en una estrategia central para modelar relaciones entre proteínas, genes, fármacos, enfermedades y funciones biológicas. En este contexto, las Graph Neural Networks (GNNs) han impulsado avances importantes porque permiten combinar estructura relacional y atributos biológicos en tareas de predicción. Sin embargo, el crecimiento de modelos no siempre ha estado acompañado por benchmarks suficientemente auditables. Iniciativas generales como Open Graph Benchmark establecieron prácticas reproducibles para aprendizaje en grafos (Hu et al., 2020), mientras que MoleculeNet y Therapeutics Data Commons organizaron tareas moleculares y terapéuticas comparables (Wu et al., 2018; Huang et al., 2022). En grafos biomédicos, OpenBioLink y OBNB ampliaron el panorama hacia link prediction y node classification sobre redes biológicas (Breit et al., 2020; Liu & Krishnan, 2024). Aun así, persiste una brecha específica: evaluar redes proteína-proteína humanas bajo controles explícitos de descarga, filtrado, solapamiento entre recursos, particionado, features y calibración.

Esta brecha es relevante porque recursos como STRING y BioGRID son ampliamente utilizados, abiertos y valiosos, pero no fueron diseñados originalmente como benchmarks cerrados. Contienen evidencia curada, cobertura desigual, redundancia y relaciones que pueden reaparecer entre fuentes. Si estos recursos se usan sin auditoría, un modelo puede parecer superior por capturar sesgos topológicos, reutilización indirecta de interacciones o decisiones de split poco controladas. Por ello, el problema no es solo entrenar GNNs más complejas, sino definir cuándo una mejora es atribuible a aprendizaje generalizable y cuándo responde a propiedades del protocolo experimental.

Este artículo presenta BioGraphBench, un benchmark reproducible para aprendizaje sobre grafos biomédicos centrado en confianza experimental. Siguiendo la lógica IMRAD, el trabajo parte de una pregunta metodológica: ¿pueden transformarse recursos biomédicos abiertos en tareas de evaluación que sean auditables desde los datos crudos hasta los resultados finales? Para responderla, se diseñó un pipeline que audita datasets, canonicaliza redes PPI, define tareas de link prediction y node classification, genera splits reproducibles, calcula features estructurales únicamente desde el grafo de entrenamiento y evalúa modelos mediante métricas de ranking, decisión y calibración.

El MVP de BioGraphBench integra tareas sobre STRING, BioGRID y OBNB. En predicción de enlaces, se evaluaron tareas principales y ablaciones sin solapamiento entre STRING y BioGRID. En clasificación de nodos, se construyó una tarea multilabel BioGRID+GOBP orientada a recuperar funciones biológicas. Los baselines incluyen heurísticas clásicas de grafos, Logistic Regression, Random Forest, HistGradientBoosting y modelos piloto MLP/GCN. Esta elección metodológica es deliberada: antes de afirmar que una GNN aporta aprendizaje biomédico adicional, debe superar señales estructurales simples pero fuertes, calculadas bajo restricciones reproducibles.

Los resultados muestran que BioGraphBench produce evidencia útil desde su primera versión. En link prediction PPI, HistGradientBoosting alcanzó AUPRC entre 0,942 y 0,963, con ganancias sobre la mejor heurística de hasta +0,042 y ECE-10 muy bajo, entre 0,0037 y 0,0060. Estos valores indican que la topología PPI contiene una señal predictiva fuerte y que futuras GNNs deberán superar baselines no neuronales exigentes, no solo heurísticas débiles. En contraste, la tarea de node classification fue sustancialmente más difícil: el GCN piloto alcanzó Macro AUPRC=0,016 y Micro-F1 cercano a 0,021, mostrando que propagar información por la red no resuelve por sí sola el desbalance funcional.

La contribución de BioGraphBench es, por tanto, metodológica y empírica. Metodológica, porque convierte la reproducibilidad en una propiedad de diseño del benchmark mediante artefactos, scripts, splits y reportes verificables. Empírica, porque establece una línea base realista para PPI y evidencia una tarea funcional todavía abierta. En conjunto, el artículo propone que un benchmark biomédico serio no debe empezar con modelos, sino con confianza: datos abiertos, tareas definidas sin leakage, baselines fuertes, métricas calibradas y resultados reconstruibles."""


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("2E74B5")

    for line in INTRODUCTION.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)

    doc.save(OUT_DOCX)


def main() -> None:
    OUT_MD.write_text(INTRODUCTION, encoding="utf-8")
    build_docx()
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
